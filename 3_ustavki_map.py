#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
3_ustavki_map.py — Редактирование карт уставок (Visio) и отчёт изменений
=========================================================================
Выполняет шаги 5–6 цикла обработки таблиц уставок:
  5. Извлечь синие строки (изменения) из таблиц и создать отчёт Word
  6. Обновить гиперссылки в картах Visio (.vsdx) + экспорт в PDF

ВХОДНЫЕ ДАННЫЕ:
  - session_data.json → ключ "ustavki": список UstavkiEntry
    (сгенерированный программой 2_ustavki_folders.py)
  - Или вручную добавить/ввести записи

ИНТЕРФЕЙС ДЛЯ СШИВАНИЯ:
  - UstavkiMapApp.ustavki_entries  — публичный список записей
  - session_data.json → ключ "ustavki" обновляется после каждого шага
  - get_entries() → list[dict]  — передать в программу 4
"""

import os
import re
import sys
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime

from shared_lib import (
    MAPS_FOLDER, MAPS_PDF_FOLDER, USTAVKI_ARCHIVE_BASE,
    EMPTY_USTAVKI_ENTRY, match_object_to_short_name,
    get_object_short_name_from_path,
    load_session, save_session,
)

# ── Зависимости ──────────────────────────────────────────────────────────────

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
    _BASE_CLASS = TkinterDnD.Tk
except Exception:
    HAS_DND = False
    _BASE_CLASS = tk.Tk

# ── Синие строки ─────────────────────────────────────────────────────────────

_BLUE_COLORS = {c.lower() for c in [
    '4472C4','5B9BD5','2E74B5','0070C0','00B0F0','1F3864',
    '2F5597','1155CC','0000FF','538DD5','4F81BD','44546A',
]}


def _run_has_blue_text(run_element) -> bool:
    """
    Проверяет, есть ли в run синий цвет И есть ли в нём буквы/цифры.
    Табуляции и пробелы с синим цветом — игнорируем (легаси форматирование).
    """
    try:
        from lxml import etree
    except ImportError:
        return False
    # Текст run — только если содержит хотя бы одну букву или цифру
    run_text = ''.join(t.text or '' for t in run_element.iter()
                       if t.tag.endswith('}t') or t.tag == 't')
    if not re.search(r'[a-zA-Zа-яА-ЯёЁ0-9]', run_text):
        return False  # Таб, пробел, спецсимволы — пропускаем
    # Проверяем цвет в XML данного run
    xml_str = etree.tostring(run_element).decode('utf-8', errors='ignore').lower()
    for bc in _BLUE_COLORS:
        if bc in xml_str:
            return True
    return False


def extract_blue_rows_from_doc(doc_path: str) -> list:
    """
    Возвращает строки таблиц, содержащие синий текст.
    Знаки табуляции с синим цветом НЕ учитываются — только буквы и цифры.
    """
    if not HAS_DOCX:
        return []
    try:
        from lxml import etree
        from docx.oxml.ns import qn as _qn
    except ImportError:
        return []
    doc = DocxDocument(doc_path)
    # Тег run в XML
    W_R = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'
    result = []
    for table in doc.tables:
        for row in table.rows:
            has_blue = False
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run_el in para._element.iter(W_R):
                        if _run_has_blue_text(run_el):
                            has_blue = True
                            break
                    if has_blue:
                        break
                if has_blue:
                    break
            if has_blue:
                result.append([c.text.strip() for c in row.cells])
    return result


def generate_changes_report(entries: list, output_path: str):
    """Создаёт Word-документ со сводкой синих строк из всех таблиц."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx не установлен")
    doc = DocxDocument()
    doc.add_heading('Сводка изменений таблиц уставок', level=1)
    for i, entry in enumerate(entries, 1):
        doc.add_heading(
            f"Таблица {i}: {entry.get('dispatch_name','?')} — {entry.get('table_number','?')}",
            level=2)
        doc.add_paragraph(
            f"Объект: {entry.get('object_name','')}\n"
            f"Исх. письмо: {entry.get('outgoing_letter','')}"
        )
        blue_rows = extract_blue_rows_from_doc(entry.get('file_path', ''))
        if not blue_rows:
            doc.add_paragraph('Изменений (синий цвет) не обнаружено.')
            continue
        num_cols = max(len(r) for r in blue_rows)
        tbl = doc.add_table(rows=len(blue_rows), cols=num_cols)
        tbl.style = 'Table Grid'
        for ri, row_data in enumerate(blue_rows):
            for ci, cell_text in enumerate(row_data):
                if ci < num_cols:
                    tbl.cell(ri, ci).text = cell_text
    doc.save(output_path)
    try:
        os.startfile(output_path)
    except Exception:
        pass


# ── Visio: обновление карт ────────────────────────────────────────────────────

def update_visio_map(visio_path: str, old_table_path: str,
                     new_table_path: str, new_table_number: str) -> tuple:
    """
    Открывает Visio, заменяет гиперссылку old→new, сохраняет и экспортирует PDF.
    Возвращает (success: bool, message: str).
    """
    try:
        import win32com.client as win32
    except ImportError:
        return False, "win32com не доступен (нужен pywin32)"
    if not os.path.exists(visio_path):
        return False, f"Файл не найден: {visio_path}"

    visio = None
    try:
        visio = win32.Dispatch('Visio.Application')
        visio.Visible = False
        doc = visio.Documents.Open(os.path.abspath(visio_path))
        replaced = 0
        for page in doc.Pages:
            for shape in page.Shapes:
                for i in range(1, shape.Hyperlinks.Count + 1):
                    hl = shape.Hyperlinks.Item(i)
                    if old_table_path.lower() in hl.Address.lower():
                        hl.Address = new_table_path
                        if new_table_number:
                            hl.Description = new_table_number
                        replaced += 1
        doc.Save()
        stem = os.path.splitext(os.path.basename(visio_path))[0]
        pdf_path = os.path.join(MAPS_PDF_FOLDER, stem + '.pdf')
        os.makedirs(MAPS_PDF_FOLDER, exist_ok=True)
        doc.ExportAsFixedFormat(1, pdf_path, 0, 0)
        doc.Close()
        return True, f"Заменено ссылок: {replaced}, PDF: {pdf_path}"
    except Exception as exc:
        return False, str(exc)
    finally:
        try:
            if visio:
                visio.Quit()
        except Exception:
            pass


# ── Главное окно ─────────────────────────────────────────────────────────────

class UstavkiMapApp(_BASE_CLASS):
    """
    Программа 3: Обновление карт уставок.

    Публичный интерфейс:
      .ustavki_entries   list[dict]  — обрабатываемые записи
      .get_entries() → list[dict]    — результат для программы 4
    """

    def __init__(self):
        super().__init__()
        self.ustavki_entries: list = []

        session = load_session()
        if 'ustavki' in session:
            self.ustavki_entries = session['ustavki']

        self.title("Таблицы уставок — Карты (Visio)  v2")
        self.resizable(True, True)
        self._build_ui()
        self._center_window()

    def get_entries(self) -> list:
        return [dict(e) for e in self.ustavki_entries]

    # ── UI ────────────────────────────────────────────────────────────────

    def _build_ui(self):
        root = ttk.Frame(self, padding=6)
        root.grid(row=0, column=0, sticky='nsew')
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        root.columnconfigure(0, weight=1)
        root.rowconfigure(0, weight=1)

        self._step_nb = ttk.Notebook(root)
        self._step_nb.grid(row=0, column=0, sticky='nsew', pady=4)

        step5 = ttk.Frame(self._step_nb, padding=8)
        step6 = ttk.Frame(self._step_nb, padding=8)
        self._step_nb.add(step5, text=" 5 Изменения (синие строки) ")
        self._step_nb.add(step6, text=" 6 Карты Visio + PDF ")

        self._logs = {}
        self._build_step5(step5)
        self._build_step6(step6)

        bot = ttk.Frame(root)
        bot.grid(row=1, column=0, pady=4, sticky='ew')
        ttk.Button(bot, text="Загрузить сессию",   command=self._load_session).pack(side='left', padx=4)
        ttk.Button(bot, text="Сохранить сессию",   command=self._save_session).pack(side='left', padx=4)
        ttk.Button(bot, text="Редактировать список", command=self._edit_entries).pack(side='left', padx=4)
        ttk.Separator(bot, orient='vertical').pack(side='left', fill='y', padx=6)
        ttk.Button(bot, text="▶ Шаг 5: Отчёт изменений",
                   command=self._create_changes_report).pack(side='left', padx=4)
        ttk.Button(bot, text="▶ Шаг 6: Обновить карты Visio",
                   command=self._update_maps_all).pack(side='left', padx=4)
        ttk.Button(bot, text="Закрыть",            command=self.destroy).pack(side='right', padx=4)

    def _build_step5(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)
        ttk.Label(parent,
            text="Ищет строки синего цвета (изменения) во всех таблицах .docx "
                 "и формирует сводный отчёт Word.\n"
                 "Файлы берутся из сессии (поле file_path каждой записи).",
            wraplength=700, foreground='gray').grid(row=0, column=0, sticky='w', pady=(0,6))

        txt = tk.Text(parent, height=14, wrap='word', state='disabled', font=('Consolas', 9))
        sb = ttk.Scrollbar(parent, orient='vertical', command=txt.yview)
        txt.configure(yscrollcommand=sb.set)
        txt.grid(row=1, column=0, sticky='nsew')
        sb.grid(row=1, column=1, sticky='ns')
        self._logs['5'] = txt

        bf = ttk.Frame(parent)
        bf.grid(row=2, column=0, pady=(6,0), sticky='w')
        ttk.Button(bf, text="Создать отчёт изменений →",
                   command=self._create_changes_report).pack(side='left', padx=4)

    def _build_step6(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(2, weight=1)

        # Настройки путей
        cfg = ttk.LabelFrame(parent, text="Пути к папкам", padding=6)
        cfg.grid(row=0, column=0, sticky='ew', pady=(0, 4))
        cfg.columnconfigure(1, weight=1)
        ttk.Label(cfg, text="Папка карт Visio:").grid(row=0, column=0, sticky='e', padx=(0, 6))
        self._maps_folder_var = tk.StringVar(value=MAPS_FOLDER)
        ttk.Entry(cfg, textvariable=self._maps_folder_var, width=50).grid(row=0, column=1, sticky='ew')
        ttk.Button(cfg, text="…", command=lambda: self._browse_folder(self._maps_folder_var),
                   width=3).grid(row=0, column=2)
        ttk.Label(cfg, text="Папка PDF:").grid(row=1, column=0, sticky='e', padx=(0, 6))
        self._pdf_folder_var = tk.StringVar(value=MAPS_PDF_FOLDER)
        ttk.Entry(cfg, textvariable=self._pdf_folder_var, width=50).grid(row=1, column=1, sticky='ew')
        ttk.Button(cfg, text="…", command=lambda: self._browse_folder(self._pdf_folder_var),
                   width=3).grid(row=1, column=2)

        # Ручной ввод данных для одной таблицы (если сессия не полная)
        inp = ttk.LabelFrame(parent,
            text="Ручной ввод (заполняется автоматически из сессии или вручную)", padding=6)
        inp.grid(row=1, column=0, sticky='ew', pady=(0, 4))
        inp.columnconfigure(1, weight=1)

        ttk.Label(inp, text="Старый файл таблицы\n(гиперссылка для замены):",
                  justify='right').grid(row=0, column=0, sticky='e', padx=(0, 6), pady=3)
        self._old_table_path_var = tk.StringVar()
        old_row = ttk.Frame(inp); old_row.grid(row=0, column=1, sticky='ew', pady=3)
        old_row.columnconfigure(0, weight=1)
        ttk.Entry(old_row, textvariable=self._old_table_path_var, width=52).grid(
            row=0, column=0, sticky='ew')
        ttk.Button(old_row, text="…",
                   command=lambda: self._browse_file(self._old_table_path_var),
                   width=3).grid(row=0, column=1, padx=(4, 0))

        ttk.Label(inp, text="Новый файл таблицы\n(будет вставлен как ссылка):",
                  justify='right').grid(row=1, column=0, sticky='e', padx=(0, 6), pady=3)
        self._new_table_path_var = tk.StringVar()
        new_row = ttk.Frame(inp); new_row.grid(row=1, column=1, sticky='ew', pady=3)
        new_row.columnconfigure(0, weight=1)
        ttk.Entry(new_row, textvariable=self._new_table_path_var, width=52).grid(
            row=0, column=0, sticky='ew')
        ttk.Button(new_row, text="…",
                   command=lambda: self._browse_file(self._new_table_path_var),
                   width=3).grid(row=0, column=1, padx=(4, 0))
        ttk.Label(inp,
            text="(при наличии сессии — новый путь берётся из current_path каждой записи)",
            foreground='gray').grid(row=2, column=0, columnspan=2, sticky='w', pady=(0, 2))

        ttk.Label(inp, text="Объект (краткое имя\n= имя папки):").grid(
            row=3, column=0, sticky='e', padx=(0, 6), pady=3)
        obj_row = ttk.Frame(inp); obj_row.grid(row=3, column=1, sticky='ew', pady=3)
        obj_row.columnconfigure(0, weight=1)
        self._manual_object_var = tk.StringVar()
        ttk.Entry(obj_row, textvariable=self._manual_object_var, width=30).grid(
            row=0, column=0, sticky='w')
        ttk.Button(obj_row, text="Определить из пути нового файла",
                   command=self._auto_detect_object).grid(row=0, column=1, padx=(8, 0))

        ttk.Label(parent,
            text="Для каждой таблицы из списка:\n"
                 "  • Находит .vsdx файл по краткому имени объекта\n"
                 "  • Заменяет гиперссылку (старый → новый путь к .docx)\n"
                 "  • Сохраняет .vsdx и экспортирует PDF\n"
                 "ТРЕБУЕТ: Microsoft Visio и pywin32",
            wraplength=700, foreground='gray').grid(row=2, column=0, sticky='w', pady=(0, 4))

        txt = tk.Text(parent, height=10, wrap='word', state='disabled', font=('Consolas', 9))
        sb = ttk.Scrollbar(parent, orient='vertical', command=txt.yview)
        txt.configure(yscrollcommand=sb.set)
        txt.grid(row=3, column=0, sticky='nsew')
        sb.grid(row=3, column=1, sticky='ns')
        self._logs['6'] = txt

        bf = ttk.Frame(parent)
        bf.grid(row=4, column=0, pady=(6, 0), sticky='w')
        ttk.Button(bf, text="Обновить карты Visio (из сессии) →",
                   command=self._update_maps_all).pack(side='left', padx=4)
        ttk.Button(bf, text="Обновить карту для одного объекта →",
                   command=self._update_single_map).pack(side='left', padx=4)

    def _browse_folder(self, var: tk.StringVar):
        d = filedialog.askdirectory(initialdir=var.get() or os.path.expanduser('~'), parent=self)
        if d:
            var.set(d)

    def _browse_file(self, var: tk.StringVar):
        initial = os.path.dirname(var.get()) if var.get() else os.path.expanduser('~')
        f = filedialog.askopenfilename(
            initialdir=initial,
            filetypes=[("Word файлы", "*.docx *.doc"), ("Все файлы", "*.*")],
            parent=self,
        )
        if f:
            var.set(f.replace('/', '\\'))

    def _auto_detect_object(self):
        """Определяет краткое имя объекта из пути нового файла."""
        new_path = self._new_table_path_var.get().strip()
        if not new_path:
            messagebox.showwarning("Нет пути", "Сначала выберите новый файл таблицы.", parent=self)
            return
        short = get_object_short_name_from_path(new_path)
        if short:
            self._manual_object_var.set(short)
        else:
            # Fallback: просто имя родительской папки
            short = os.path.basename(os.path.dirname(new_path))
            self._manual_object_var.set(short)
            messagebox.showinfo("Объект", f"Объект не найден в справочнике.\n"
                                          f"Установлено: {short}", parent=self)

    # ── Логирование ───────────────────────────────────────────────────────

    def _log(self, key: str, text: str):
        widget = self._logs.get(key)
        if widget:
            widget.configure(state='normal')
            widget.insert('end', text + '\n')
            widget.see('end')
            widget.configure(state='disabled')

    # ── Шаг 5: изменения ─────────────────────────────────────────────────

    def _create_changes_report(self):
        if not HAS_DOCX:
            self._log('5', "ОШИБКА: python-docx не установлен"); return
        if not self.ustavki_entries:
            self._log('5', "Список записей пуст. Загрузите сессию."); return

        default_name = f"Изменения_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        out_path = filedialog.asksaveasfilename(
            title="Сохранить отчёт изменений",
            initialdir=os.path.expanduser('~'),
            initialfile=default_name,
            defaultextension='.docx',
            filetypes=[("Word файлы", "*.docx"), ("Все файлы", "*.*")],
            parent=self,
        )
        if not out_path:
            return

        # Подтверждение
        msg = (f"Создать отчёт изменений для {len(self.ustavki_entries)} таблиц?\n\n"
               f"Сохранить как:\n{out_path}\n\nПродолжить?")
        if not messagebox.askyesno("Подтверждение", msg, parent=self):
            return

        self._log('5', f"Анализ {len(self.ustavki_entries)} файлов...")
        try:
            generate_changes_report(self.ustavki_entries, out_path)
            self._log('5', f"Отчёт создан: {out_path}")
        except Exception as exc:
            self._log('5', f"ОШИБКА: {exc}")

    # ── Шаг 6: карты Visio ───────────────────────────────────────────────

    def _update_maps_all(self):
        maps_folder = self._maps_folder_var.get()
        if not self.ustavki_entries:
            self._log('6', "Список записей пуст. Загрузите сессию."); return

        # Подтверждение
        objects_to_update = []
        for entry in self.ustavki_entries:
            short = match_object_to_short_name(entry.get('object_name', ''))
            if short:
                objects_to_update.append(short)

        msg = (f"Обновить карты Visio для {len(objects_to_update)} объектов?\n\n"
               f"Папка карт: {maps_folder}\n\n"
               + '\n'.join(f"  • {o}" for o in objects_to_update[:10])
               + (f"\n  ...и ещё {len(objects_to_update)-10}" if len(objects_to_update) > 10 else '')
               + "\n\nЭто откроет Microsoft Visio для каждого объекта!\nПродолжить?")
        if not messagebox.askyesno("Подтверждение", msg, parent=self):
            return

        for entry in self.ustavki_entries:
            short = match_object_to_short_name(entry.get('object_name', ''))
            if not short:
                self._log('6', f"Объект не распознан: {entry.get('object_name','')}")
                continue

            visio_path = os.path.join(maps_folder, short + '.vsdx')
            if not os.path.exists(visio_path):
                visio_path = os.path.join(maps_folder, short + '.vsd')
            if not os.path.exists(visio_path):
                self._log('6', f"Visio не найден: {short}.vsdx / .vsd")
                continue

            old_path   = entry.get('archive_candidate', '')
            new_path   = entry.get('current_path', entry.get('file_path', ''))
            table_num  = entry.get('table_number', '')

            self._log('6', f"Обновляю: {short}  {os.path.basename(visio_path)}")
            ok, msg_result = update_visio_map(visio_path, old_path, new_path, table_num)
            self._log('6', f"  {'OK' if ok else 'ERR'}  {msg_result}")

        self._save_session()
        self._log('6', "--- Завершено ---")

    def _update_single_map(self):
        """Обновляет карту для одного объекта по данным из формы ввода."""
        old_path  = self._old_table_path_var.get().strip()
        new_path  = self._new_table_path_var.get().strip()
        short     = self._manual_object_var.get().strip()

        if not new_path:
            messagebox.showwarning("Нет данных", "Укажите новый файл таблицы.", parent=self)
            return
        if not short:
            messagebox.showwarning("Нет данных", "Укажите краткое имя объекта (или нажмите «Определить»).", parent=self)
            return

        maps_folder = self._maps_folder_var.get()
        visio_path = os.path.join(maps_folder, short + '.vsdx')
        if not os.path.exists(visio_path):
            visio_path = os.path.join(maps_folder, short + '.vsd')
        if not os.path.exists(visio_path):
            self._log('6', f"Visio не найден: {short}.vsdx / .vsd в {maps_folder}")
            return

        msg = (f"Обновить карту?\n\n"
               f"  Объект:    {short}\n"
               f"  Visio:     {os.path.basename(visio_path)}\n"
               f"  Старый:    {os.path.basename(old_path) if old_path else '(не указан)'}\n"
               f"  Новый:     {os.path.basename(new_path)}\n\n"
               f"Откроется Microsoft Visio. Продолжить?")
        if not messagebox.askyesno("Подтверждение", msg, parent=self):
            return

        self._log('6', f"Обновляю: {short}  {os.path.basename(visio_path)}")
        ok, result_msg = update_visio_map(visio_path, old_path, new_path, '')
        self._log('6', f"  {'OK' if ok else 'ERR'}  {result_msg}")
        self._save_session()

    # ── Редактирование записей ────────────────────────────────────────────

    def _edit_entries(self):
        """Простой диалог для ввода/редактирования списка file_path."""
        dlg = tk.Toplevel(self)
        dlg.title("Записи (file_path и данные)")
        dlg.grab_set()
        dlg.columnconfigure(0, weight=1)
        dlg.rowconfigure(0, weight=1)

        cols = ('file', 'object', 'dispatch', 'table_num', 'status')
        tv = ttk.Treeview(dlg, columns=cols, show='headings', height=12)
        tv.heading('file',      text='Файл');        tv.column('file',      width=200)
        tv.heading('object',    text='Объект');       tv.column('object',    width=150)
        tv.heading('dispatch',  text='Дисп. наим.'); tv.column('dispatch',  width=200)
        tv.heading('table_num', text='№ таблицы');   tv.column('table_num', width=90)
        tv.heading('status',    text='Статус');       tv.column('status',    width=100)
        sb = ttk.Scrollbar(dlg, orient='vertical', command=tv.yview)
        tv.configure(yscrollcommand=sb.set)
        tv.grid(row=0, column=0, sticky='nsew', padx=6, pady=6)
        sb.grid(row=0, column=1, sticky='ns')

        for entry in self.ustavki_entries:
            tv.insert('', 'end', values=(
                os.path.basename(entry.get('file_path','')),
                entry.get('object_name',''),
                entry.get('dispatch_name',''),
                entry.get('table_number',''),
                entry.get('status',''),
            ))

        bf = ttk.Frame(dlg)
        bf.grid(row=1, column=0, pady=6)
        ttk.Button(bf, text="Добавить файлы…", command=lambda: self._add_files_dlg(tv)).pack(side='left', padx=4)
        ttk.Button(bf, text="Закрыть",         command=dlg.destroy).pack(side='right', padx=4)
        dlg.transient(self)

    def _add_files_dlg(self, tv):
        files = filedialog.askopenfilenames(
            title="Выберите .docx",
            filetypes=[("Word файлы", "*.docx *.doc"), ("Все файлы", "*.*")],
        )
        for f in files:
            for e in self.ustavki_entries:
                if e['file_path'] == f:
                    break
            else:
                entry = dict(EMPTY_USTAVKI_ENTRY)
                entry['file_path'] = f
                self.ustavki_entries.append(entry)
                tv.insert('', 'end', values=(os.path.basename(f), '', '', '', 'добавлено'))

    # ── Сессия ────────────────────────────────────────────────────────────

    def _load_session(self):
        session = load_session()
        if 'ustavki' in session:
            self.ustavki_entries = session['ustavki']
            messagebox.showinfo("Сессия",
                f"Загружено {len(self.ustavki_entries)} записей.", parent=self)
        else:
            messagebox.showinfo("Сессия", "Данные уставок не найдены в сессии.", parent=self)

    def _save_session(self):
        session = load_session()
        session['ustavki'] = self.ustavki_entries
        save_session(session)

    # ── Центровка ─────────────────────────────────────────────────────────

    def _center_window(self):
        self.update_idletasks()
        w, h = self.winfo_width(), self.winfo_height()
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        self.geometry(f"{max(w,900)}x{max(h,600)}+{(sw-max(w,900))//2}+{(sh-max(h,600))//2}")


# ── Точка входа ───────────────────────────────────────────────────────────────

def main():
    app = UstavkiMapApp()
    app.mainloop()


if __name__ == '__main__':
    main()
