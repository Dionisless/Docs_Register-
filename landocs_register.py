#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Регистратор корреспонденции LanDocs v2
========================================
Модуль 1: Регистрация входящих/исходящих писем в журнал Excel.
Модуль 2: Обработка таблиц уставок релейной защиты (7 шагов).

Позиции полей ВХОДЯЩИХ (Tab от начала):
  0  — № вх
  3  — ссылка на файл письма
  4  — Корреспондент
  5  — Дата
  6  — № письма
  8  — Подписант
  10 — Тема письма
  15 — Связанное письмо

Позиции полей ИСХОДЯЩИХ (Tab от начала):
  0  — № письма
  1  — Дата
  5  — Тема письма
  6  — Исполнитель
  9  — ФИО получателей (через ;)
  10 — Компании получателей (через ;)
  16 — Связанное письмо
"""

import os
import re
import sys
import time
import shutil
import getpass
import difflib
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime

# ── Зависимости с мягкой обработкой ─────────────────────────────────────────

try:
    import win32api, win32con, win32clipboard
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

try:
    import openpyxl
    from openpyxl.styles import Alignment
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
    _BASE_CLASS = TkinterDnD.Tk
except Exception:
    HAS_DND = False
    _BASE_CLASS = tk.Tk

try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.options import Options as ChromeOptions
    HAS_SELENIUM = True
except ImportError:
    HAS_SELENIUM = False

try:
    import pyautogui
    HAS_PYAUTOGUI = True
except ImportError:
    HAS_PYAUTOGUI = False

# ── Конфигурация ─────────────────────────────────────────────────────────────

EXCEL_PATH_IN  = r"\\Prim-fs-serv\primrdu\СРЗА\Журнал регистрации корреспонденции\Журнал регистрации входящей документации.xlsx"
EXCEL_PATH_OUT = r"\\Prim-fs-serv\primrdu\СРЗА\Журнал регистрации корреспонденции\Журнал регистрации исходящей документации.xlsx"
DEFAULT_SAVE_FOLDER = r"\\Prim-fs-serv\primrdu\СРЗА\Дела СРЗА\19 Переписка"

REGISTRY_PATH    = r"\\Prim-fs-serv\primrdu\СРЗА\Журнал регистрации корреспонденции\Реестр таблиц уставок.xlsx"
SUMMARY_PATH     = r"\\Prim-fs-serv\primrdu\СРЗА\Журнал регистрации корреспонденции\Регистрация таблиц уставок.xlsx"
USTAVKI_EXEC_BASE = r"\\Prim-fs-serv\primrdu\СРЗА\Уставки\Таблицы РЗА\Таблицы для исполнения РЗА"
MAPS_FOLDER      = r"\\Prim-fs-serv\primrdu\СРЗА\Уставки\КАРТА УСТАВОК"
MAPS_PDF_FOLDER  = r"\\Prim-fs-serv\primrdu\СРЗА\Уставки\КАРТА УСТАВОК\ДЭБ"
DEB_BASE_URL     = "https://pri-mdeb.oduvs.so"
DEB_MAPS_URL     = "https://pri-mdeb.oduvs.so/?sid=02ab815f-a54e-42a5-8a88-36dee8a5af2e&DataAreaId=1b6fecd6-f813-47ac-aa88-de4f67b7a1ac"

# Задержки (сек)
TAB_DELAY  = 0.07
COPY_DELAY = 0.12

# Сокращённые названия объектов (жаргонные, = имена папок)
OBJECT_SHORT_NAMES = [
    "ПримГРЭС","ТЭЦ Центральная","АТЭЦ","ВТЭЦ-2","ПартГРЭС","Шкотовская ТЭЦ",
    "Восточная ТЭЦ","ТЭЦ Северная","Варяг","Владивосток","Дальневосточная","Лозовая",
    "Хехцир 2","Чугуевка 2","Арсеньев 2","Аэропорт","Береговая 2","Бикин тяговая",
    "Волна","Высокогорск","Горелое","Губерово тяговая","Западная","Звезда",
    "Зелёный угол","Иман","К","Кировка","Козьмино","Лесозаводск","Минеральная",
    "НПС-36","НПС-38","НПС-40","НПС-41","Находка","Новая","Партизанск","Патрокл",
    "Перевал","Промпарк","Розенгартовка тяговая","Ружино тяговая","Русская",
    "Свиягино тяговая","Спасск","Суходол","Уссурийск 2","Чугуевка","Широкая",
    "Шмаковка тяговая","178Ф","1Р","1Р тяговая","2Р","А","АСБ","Агрокомплекс",
    "Амурская","Анисимовка тяговая","Арсеньев 1","Барабаш","Береговая 1","Бикин",
    "Богополь","Бурная","Бурун","ВТЭЦ-1","Вадимовка","Вокзальная тяговая","Волчанец",
    "Восток 2","Восточная тяговая","Глубинная","Голдобин","Голубинка","Голубовка",
    "Горбуша","Горностай","Гранит","Давыдовка","Дальнереченск тяговая","Де-Фриз",
    "Дмитриевка","Екатериновка","ЖБИ-130","ЖБФ","З","Загородная","Залив","Игнатьевка",
    "Казармы","Камень-Рыболов","Кипарисово","Ключи","Кожзавод","Котельная 2Р",
    "Краскино","Кролевцы","ЛРЗ","Лазурная","Ласточка тяговая","Липовцы","Лучегорск",
    "М","Междуречье","Мингородок","Михайловка","Молодежная","Муравейка","Мучная",
    "Мыс Астафьева","НСРЗ","Надаровская","Надеждинская тяговая","Насосная",
    "Находка 110","Находка тяговая","Николаевка","Новицкое","Новоникольск",
    "Новопокровка","Новый мир","Океан","Ольга","Орлиная","Павловка 1","Павловка 2",
    "Песчаная","Петровичи","Плавзавод","Пластун","Подъяпольск","Полевая",
    "Преображение","Прибой","Приозерная","Прогресс","Промузел","Промысловка",
    "Промышленная","Пушкинская","Раздольное 1","Раздольное 2","Разрез","Ракушка",
    "Реттиховка","Рощино","С-55","Садовая","Седанка","Сибирцево тяговая","Славянка",
    "Смоляниново тяговая","Спасск тяговая","Спутник","Стройиндустрия","Студгородок",
    "Тайфун","Тереховка","Тимофеевка","Топаз","Троица","УКФ","Угольная","Улисс",
    "Уссурийск 1","Уссурийск тяговая","Учебная","Факел","Фридман тяговая","ХФЗ",
    "Хороль","Чайка","Черемшаны","Черниговка","Чуркин","Шахта 7","Штыково","Южная",
    "Ярославка",
]


# ── Клавиатура / буфер обмена ────────────────────────────────────────────────

def _clear_clipboard():
    if not HAS_WIN32:
        return
    try:
        win32clipboard.OpenClipboard()
        win32clipboard.EmptyClipboard()
    except Exception:
        pass
    finally:
        try:
            win32clipboard.CloseClipboard()
        except Exception:
            pass


def _get_clipboard() -> str:
    if not HAS_WIN32:
        return ""
    try:
        win32clipboard.OpenClipboard()
        try:
            if win32clipboard.IsClipboardFormatAvailable(win32con.CF_UNICODETEXT):
                data = win32clipboard.GetClipboardData(win32con.CF_UNICODETEXT)
                return (data or "").strip()
        finally:
            win32clipboard.CloseClipboard()
    except Exception:
        pass
    return ""


def _send_tab():
    if not HAS_WIN32: return
    win32api.keybd_event(win32con.VK_TAB, 0, 0, 0)
    win32api.keybd_event(win32con.VK_TAB, 0, win32con.KEYEVENTF_KEYUP, 0)


def _send_shift_tab():
    if not HAS_WIN32: return
    win32api.keybd_event(win32con.VK_SHIFT, 0, 0, 0)
    win32api.keybd_event(win32con.VK_TAB, 0, 0, 0)
    win32api.keybd_event(win32con.VK_TAB, 0, win32con.KEYEVENTF_KEYUP, 0)
    win32api.keybd_event(win32con.VK_SHIFT, 0, win32con.KEYEVENTF_KEYUP, 0)


def _send_ctrl_a():
    if not HAS_WIN32: return
    win32api.keybd_event(win32con.VK_CONTROL, 0, 0, 0)
    win32api.keybd_event(ord('A'), 0, 0, 0)
    win32api.keybd_event(ord('A'), 0, win32con.KEYEVENTF_KEYUP, 0)
    win32api.keybd_event(win32con.VK_CONTROL, 0, win32con.KEYEVENTF_KEYUP, 0)


def _send_ctrl_c():
    if not HAS_WIN32: return
    win32api.keybd_event(win32con.VK_CONTROL, 0, 0, 0)
    win32api.keybd_event(ord('C'), 0, 0, 0)
    win32api.keybd_event(ord('C'), 0, win32con.KEYEVENTF_KEYUP, 0)
    win32api.keybd_event(win32con.VK_CONTROL, 0, win32con.KEYEVENTF_KEYUP, 0)


def read_current_field() -> str:
    _clear_clipboard()
    time.sleep(0.04)
    _send_ctrl_a()
    time.sleep(0.04)
    _send_ctrl_c()
    time.sleep(COPY_DELAY)
    return _get_clipboard()


def navigate_tabs(n: int):
    for _ in range(n):
        _send_tab()
        time.sleep(TAB_DELAY)


# ── Парсинг LanDocs ──────────────────────────────────────────────────────────

def extract_landocs_data_in() -> dict:
    data = {}
    current = 0
    _send_tab(); time.sleep(TAB_DELAY)
    _send_shift_tab(); time.sleep(TAB_DELAY)
    data['incoming_num'] = read_current_field()
    navigate_tabs(3 - current); current = 3
    data['file_link'] = read_current_field()
    navigate_tabs(4 - current); current = 4
    data['correspondent'] = read_current_field()
    navigate_tabs(5 - current); current = 5
    data['date'] = read_current_field()
    navigate_tabs(6 - current); current = 6
    data['letter_num'] = read_current_field()
    navigate_tabs(8 - current); current = 8
    data['signatory'] = read_current_field()
    navigate_tabs(10 - current); current = 10
    data['subject'] = read_current_field()
    navigate_tabs(15 - current)
    data['related'] = read_current_field()
    return data


def extract_landocs_data_out() -> dict:
    data = {}
    current = 0
    _send_tab(); time.sleep(TAB_DELAY)
    _send_shift_tab(); time.sleep(TAB_DELAY)
    data['letter_num'] = read_current_field()
    navigate_tabs(1 - current); current = 1
    data['date'] = read_current_field()
    navigate_tabs(5 - current); current = 5
    data['subject'] = read_current_field()
    navigate_tabs(6 - current); current = 6
    data['executor'] = read_current_field()
    navigate_tabs(9 - current); current = 9
    data['recipient_names'] = read_current_field()
    navigate_tabs(10 - current); current = 10
    data['recipient_companies'] = read_current_field()
    navigate_tabs(16 - current)
    data['related'] = read_current_field()
    return data


# ── Утилиты ──────────────────────────────────────────────────────────────────

def build_recipient_string(names_str: str, companies_str: str) -> str:
    names     = [n.strip() for n in names_str.split(';') if n.strip()]
    companies = [c.strip() for c in companies_str.split(';') if c.strip()]
    parts = []
    for i, (n, c) in enumerate(zip(names, companies), 1):
        parts.append(f"{i}. {n} {c}")
    return ';\n'.join(parts) if parts else names_str


def find_latest_in_viewdir() -> str:
    local_app = os.environ.get('LOCALAPPDATA') or os.path.join(
        os.environ.get('USERPROFILE', ''), 'AppData', 'Local')
    view_dir = os.path.join(local_app, 'Temp', 'ViewDir')
    if not os.path.isdir(view_dir):
        return ''
    latest_path, latest_mtime = '', 0.0
    for dirpath, _, filenames in os.walk(view_dir):
        for fname in filenames:
            fpath = os.path.join(dirpath, fname)
            try:
                mtime = os.path.getmtime(fpath)
                if mtime > latest_mtime:
                    latest_mtime, latest_path = mtime, fpath
            except OSError:
                pass
    return latest_path


def sanitize_for_filename(text: str) -> str:
    return re.sub(r'[<>:"/\\|?*\r\n\t]', '_', text)


def parse_date(date_str: str):
    for fmt in ('%d.%m.%Y', '%Y-%m-%d', '%d/%m/%Y', '%d-%m-%Y', '%Y.%m.%d'):
        try:
            return datetime.strptime(date_str.strip(), fmt)
        except ValueError:
            continue
    return None


def fmt_date_ymd(date_str: str) -> str:
    dt = parse_date(date_str)
    return dt.strftime('%Y-%m-%d') if dt else date_str


def fmt_date_dmy(date_str: str) -> str:
    dt = parse_date(date_str)
    return dt.strftime('%d.%m.%Y') if dt else date_str


def fmt_date_dmy_underscore(date_str: str) -> str:
    dt = parse_date(date_str)
    return dt.strftime('%d_%m_%Y') if dt else re.sub(r'[.\-/]', '_', date_str)


def fmt_date_ymd_underscore(date_str: str) -> str:
    dt = parse_date(date_str)
    return dt.strftime('%Y_%m_%d') if dt else date_str


def build_default_filename_in(date_str: str, incoming_num: str, letter_num: str) -> str:
    date_ymd = fmt_date_ymd(date_str)
    date_dmy = fmt_date_dmy_underscore(date_str)
    letter_clean = sanitize_for_filename(letter_num)
    return f"{date_ymd} {incoming_num}_{letter_clean}_{date_dmy}"


def build_default_filename_out(date_str: str, letter_num: str) -> str:
    date_ymd = fmt_date_ymd(date_str)
    date_dmy = fmt_date_dmy_underscore(date_str)
    letter_clean = sanitize_for_filename(letter_num)
    return f"{date_ymd} {letter_clean}_{date_dmy}"


def calc_folder_num(full_path: str) -> str:
    base = DEFAULT_SAVE_FOLDER.rstrip('\\/')
    norm = full_path.replace('/', '\\')
    if norm.lower().startswith(base.lower()):
        return norm[len(base):].lstrip('\\/')
    return norm


# ── Поиск в журнале входящих (F1) ────────────────────────────────────────────

def lookup_incoming_journal(incoming_num: str, year: str) -> dict | None:
    """
    Ищет строку в журнале входящих по № вх (col 2) на листе года.
    Возвращает dict с полями письма или None.
    """
    if not HAS_OPENPYXL:
        raise RuntimeError("openpyxl не установлен")
    if not os.path.exists(EXCEL_PATH_IN):
        raise FileNotFoundError(f"Журнал не найден:\n{EXCEL_PATH_IN}")

    wb = openpyxl.load_workbook(EXCEL_PATH_IN, read_only=True, data_only=True)

    # Выбор листа: совпадение с годом или последний рабочий лист
    service = {'служебный', 'шаблон', 'template'}
    year_sheets = [s for s in wb.sheetnames if s.lower() not in service]
    ws = None
    for s in wb.sheetnames:
        if s.strip() == year.strip():
            ws = wb[s]; break
    if ws is None:
        ws = wb[year_sheets[-1]] if year_sheets else wb.worksheets[-1]

    num_clean = incoming_num.strip().lower()
    for row in range(ws.max_row, 1, -1):
        cell_val = ws.cell(row=row, column=2).value
        if cell_val is not None and str(cell_val).strip().lower() == num_clean:
            # Гиперссылка из col 3
            hyperlink = ''
            try:
                hl = ws.cell(row=row, column=3).hyperlink
                if hl:
                    hyperlink = hl.target if hasattr(hl, 'target') else str(hl)
            except Exception:
                pass

            def _str(v):
                if v is None: return ''
                if hasattr(v, 'strftime'): return v.strftime('%d.%m.%Y')
                return str(v).strip()

            return {
                'date':         _str(ws.cell(row, 1).value),
                'incoming_num': _str(ws.cell(row, 2).value),
                'letter_num':   _str(ws.cell(row, 3).value),
                'subject':      _str(ws.cell(row, 4).value),
                'author':       _str(ws.cell(row, 5).value),
                'signed_by':    _str(ws.cell(row, 6).value),
                'folder_num':   _str(ws.cell(row, 7).value),
                'keywords':     _str(ws.cell(row, 9).value),
                'related':      _str(ws.cell(row, 10).value),
                'hyperlink':    hyperlink,
            }
    return None


# ── Запись в журнал входящих ─────────────────────────────────────────────────

def write_to_excel_in(row_data: dict):
    if not HAS_OPENPYXL:
        raise RuntimeError("openpyxl не установлен")
    if not os.path.exists(EXCEL_PATH_IN):
        raise FileNotFoundError(f"Файл журнала не найден:\n{EXCEL_PATH_IN}")
    wb = openpyxl.load_workbook(EXCEL_PATH_IN)
    ws = wb.worksheets[-1]
    last_row = ws.max_row
    while last_row > 1 and ws.cell(row=last_row, column=1).value is None:
        last_row -= 1
    new_row = last_row + 1
    ws.cell(new_row, 1).value = row_data['date']
    ws.cell(new_row, 2).value = row_data['incoming_num']
    cell_letter = ws.cell(new_row, 3)
    cell_letter.value = row_data['letter_num']
    if row_data.get('hyperlink_path'):
        cell_letter.hyperlink = row_data['hyperlink_path']
        cell_letter.style = 'Hyperlink'
    ws.cell(new_row, 4).value = row_data['subject']
    ws.cell(new_row, 5).value = row_data['author']
    cell_signed = ws.cell(new_row, 6)
    cell_signed.value = row_data['signed_by']
    cell_signed.alignment = Alignment(wrap_text=True)
    ws.cell(new_row, 7).value = row_data['folder_num']
    ws.cell(new_row, 8).value = row_data['who_registered']
    ws.cell(new_row, 9).value = row_data['keywords']
    ws.cell(new_row, 10).value = row_data['related']
    wb.save(EXCEL_PATH_IN)


def write_to_excel_out(row_data: dict):
    if not HAS_OPENPYXL:
        raise RuntimeError("openpyxl не установлен")
    if not os.path.exists(EXCEL_PATH_OUT):
        raise FileNotFoundError(f"Файл журнала не найден:\n{EXCEL_PATH_OUT}")
    wb = openpyxl.load_workbook(EXCEL_PATH_OUT)
    ws = wb.worksheets[-1]
    last_row = ws.max_row
    while last_row > 1 and ws.cell(row=last_row, column=1).value is None:
        last_row -= 1
    new_row = last_row + 1
    ws.cell(new_row, 1).value = row_data['date']
    cell_letter = ws.cell(new_row, 2)
    cell_letter.value = row_data['letter_num']
    if row_data.get('hyperlink_path'):
        cell_letter.hyperlink = row_data['hyperlink_path']
        cell_letter.style = 'Hyperlink'
    ws.cell(new_row, 3).value = row_data['subject']
    cell_recip = ws.cell(new_row, 4)
    cell_recip.value = row_data['recipient']
    cell_recip.alignment = Alignment(wrap_text=True)
    ws.cell(new_row, 5).value = row_data['executor']
    ws.cell(new_row, 6).value = row_data['keywords']
    ws.cell(new_row, 7).value = row_data['related']
    ws.cell(new_row, 8).value = row_data['control']
    wb.save(EXCEL_PATH_OUT)


# ── Word: парсинг таблиц уставок ─────────────────────────────────────────────

def detect_table_form(doc_path: str) -> str:
    """Определяет форму: 'old' (3 колонки) или 'new' (2 колонки)."""
    if not HAS_DOCX:
        raise RuntimeError("python-docx не установлен")
    doc = DocxDocument(doc_path)
    if not doc.tables:
        return 'unknown'
    t = doc.tables[0]
    if not t.rows:
        return 'unknown'
    # Старая форма: 3 колонки, первая ячейка — число
    first_cell = t.rows[0].cells[0].text.strip()
    if len(t.columns) >= 3 and first_cell.isdigit():
        return 'old'
    return 'new'


def parse_ustavki_table(doc_path: str) -> dict:
    """
    Парсит данные из таблицы уставок (Word).
    Возвращает: object_name, dispatch_name, table_number, outgoing_letter, form_type.
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx не установлен")
    doc = DocxDocument(doc_path)
    fname = os.path.basename(doc_path)
    form = detect_table_form(doc_path)
    result = {
        'form_type': form,
        'object_name': '',
        'dispatch_name': '',
        'table_number': '',
        'outgoing_letter': '',
    }

    if form == 'old':
        # Номер таблицы из заголовочного параграфа: "...ПРДУ-РЗ-25-379"
        if doc.paragraphs:
            m = re.search(r'ПРДУ-РЗ-(\d{2}-\d+)', doc.paragraphs[0].text)
            if m:
                result['table_number'] = m.group(1)
        # Объект и дисп. наим. из первой таблицы (3 колонки)
        t = doc.tables[0]
        if len(t.rows) > 0:
            result['object_name'] = t.rows[0].cells[2].text.strip()
        if len(t.rows) > 1:
            result['dispatch_name'] = t.rows[1].cells[2].text.strip()
        # Исходящее письмо из параграфа "Уставки выданы:"
        for p in doc.paragraphs:
            if 'уставки выданы' in p.text.lower():
                # "Уставки выданы: Р45-б1-V-19-1950 от 21.10.2025"
                after_colon = p.text.split(':', 1)[-1].strip()
                # Берём только номер письма (до пробела+от)
                m2 = re.match(r'([\S]+)', after_colon)
                if m2:
                    result['outgoing_letter'] = m2.group(1)
                break

    else:  # new form
        # Номер таблицы из имени файла: "...-26-25("
        m = re.search(r'-(\d{2}-\d+)\(', fname)
        if m:
            result['table_number'] = m.group(1)
        # Объект и дисп. наим. из первой таблицы (2 колонки)
        if doc.tables:
            t = doc.tables[0]
            if len(t.rows) > 0:
                result['object_name'] = t.rows[0].cells[1].text.strip()
            if len(t.rows) > 1:
                result['dispatch_name'] = t.rows[1].cells[1].text.strip()
        # Исходящее письмо из параграфа "к письму ... № LETTER"
        for p in doc.paragraphs:
            txt = p.text
            if 'к письму' in txt.lower() or re.search(r'№\s*[РрP]45', txt):
                m2 = re.search(r'№\s*([\w\-\.\/]+)\s*$', txt.strip())
                if m2:
                    result['outgoing_letter'] = m2.group(1)
                    break

    return result


# ── Word: запись "Уставки выставлены" ────────────────────────────────────────

def _levenshtein(s1: str, s2: str) -> int:
    if len(s1) < len(s2):
        return _levenshtein(s2, s1)
    if not s2:
        return len(s1)
    prev = list(range(len(s2) + 1))
    for i, c1 in enumerate(s1):
        curr = [i + 1]
        for j, c2 in enumerate(s2):
            curr.append(min(prev[j + 1] + 1, curr[j] + 1, prev[j] + (c1 != c2)))
        prev = curr
    return prev[-1]


def _find_issued_paragraph(doc):
    """
    Ищет параграф с 'уставки выставлены' (точно) или 'таблицы выданы' (≤2 ошибки).
    Возвращает (paragraph, idx) или (None, -1).
    """
    target_exact = 'уставки выставлены'
    target_fuzzy = 'таблицы выданы'
    for idx, p in enumerate(doc.paragraphs):
        low = p.text.lower()
        if target_exact in low:
            return p, idx
    # Fuzzy fallback по скользящему окну
    win_len = len(target_fuzzy)
    for idx, p in enumerate(doc.paragraphs):
        low = p.text.lower()
        for i in range(len(low) - win_len + 1):
            if _levenshtein(low[i:i + win_len], target_fuzzy) <= 2:
                return p, idx
    return None, -1


def _add_hyperlink_run(paragraph, url: str, text: str):
    """Добавляет run-гиперссылку в конец параграфа."""
    try:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True,
        )
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('r:id'), r_id)
        run_el = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        u_el = OxmlElement('w:u'); u_el.set(qn('w:val'), 'single')
        color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), '0563C1')
        rpr.append(u_el); rpr.append(color_el)
        run_el.append(rpr)
        t_el = OxmlElement('w:t'); t_el.text = text
        run_el.append(t_el)
        hl.append(run_el)
        paragraph._element.append(hl)
    except Exception:
        # Если не вышло — просто добавляем обычный run
        run = paragraph.add_run(text)
        run.font.color.rgb = None


def write_issued_to_doc(doc_path: str, incoming_letter_num: str,
                        incoming_num: str, letter_date_str: str,
                        hyperlink_path: str = '') -> bool:
    """
    Находит 'Уставки выставлены:' и дописывает туда входящее письмо.
    Формат: " НОМЕР вх-ВХ от ГГГГ_ММ_ДД"
    Возвращает True если параграф найден и изменён.
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx не установлен")
    doc = DocxDocument(doc_path)
    para, idx = _find_issued_paragraph(doc)
    if para is None:
        return False

    date_str = fmt_date_ymd_underscore(letter_date_str)
    text_to_add = f" {incoming_letter_num} вх-{incoming_num} от {date_str}"

    if hyperlink_path and os.path.exists(hyperlink_path):
        # Добавляем обычный пробел-run, потом гиперссылку
        para.add_run(f" ")
        _add_hyperlink_run(para, hyperlink_path, f"{incoming_letter_num} вх-{incoming_num} от {date_str}")
    else:
        para.add_run(text_to_add)

    doc.save(doc_path)
    return True


# ── Word: извлечение синих строк ─────────────────────────────────────────────

_BLUE_COLORS = {c.lower() for c in [
    '4472C4','5B9BD5','2E74B5','0070C0','00B0F0','1F3864',
    '2F5597','1155CC','0000FF','538DD5','4F81BD','44546A',
]}


def extract_blue_rows_from_doc(doc_path: str) -> list:
    """
    Возвращает список строк таблиц, содержащих синий текст.
    Каждая строка — list[str] текстов ячеек.
    """
    if not HAS_DOCX:
        return []
    from lxml import etree
    doc = DocxDocument(doc_path)
    result = []
    for table in doc.tables:
        for row in table.rows:
            has_blue = False
            for cell in row.cells:
                xml_bytes = etree.tostring(cell._element)
                xml_str = xml_bytes.decode('utf-8', errors='ignore').lower()
                for bc in _BLUE_COLORS:
                    if bc in xml_str:
                        has_blue = True
                        break
                if has_blue:
                    break
            if has_blue:
                result.append([c.text.strip() for c in row.cells])
    return result


def generate_changes_report(entries: list, output_path: str):
    """
    Создаёт Word-документ со сводкой синих строк из всех таблиц.
    entries: list of dict с ключами object_name, dispatch_name, file_path
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx не установлен")
    doc = DocxDocument()
    doc.add_heading('Сводка изменений таблиц уставок', level=1)
    for i, entry in enumerate(entries, 1):
        doc.add_heading(f"Таблица {i}", level=2)
        doc.add_paragraph(
            f"Объект: {entry.get('object_name','')}\t"
            f"Диспетчерское наименование: {entry.get('dispatch_name','')}"
        )
        blue_rows = extract_blue_rows_from_doc(entry.get('file_path', ''))
        if not blue_rows:
            doc.add_paragraph('Изменений (синий цвет) не обнаружено.')
            continue
        num_cols = max(len(r) for r in blue_rows)
        tbl = doc.add_table(rows=len(blue_rows), cols=num_cols)
        tbl.style = 'Table Grid'
        for ri, row_data in enumerate(blue_rows):
            for ci, cell_text in enumerate(row_data):
                tbl.cell(ri, ci).text = cell_text
    doc.save(output_path)
    try:
        os.startfile(output_path)
    except Exception:
        pass


# ── Реестр таблиц уставок (новая форма) ──────────────────────────────────────

def _norm(v) -> str:
    """Нормализация строки для сравнения."""
    if v is None: return ''
    return str(v).strip().lower()


def write_to_registry_new(entry: dict, incoming_letter_num: str,
                          incoming_num: str, letter_date) -> tuple:
    """
    Ищет строку в Реестре и записывает данные о входящем письме.
    Реестр: sheet='Реестр', заголовки в строке 2, данные с строки 3.
    col2=Объект, col3=Дисп.наим., col4=Исх.письмо
    col9=WRITE входящий №, col10=WRITE вх LanDocs, col11=WRITE дата

    Возвращает (found_row: int|None, candidates: list[dict])
    """
    if not HAS_OPENPYXL:
        raise RuntimeError("openpyxl не установлен")
    if not os.path.exists(REGISTRY_PATH):
        raise FileNotFoundError(f"Реестр не найден:\n{REGISTRY_PATH}")

    wb = openpyxl.load_workbook(REGISTRY_PATH)
    ws = wb['Реестр']

    obj   = _norm(entry.get('object_name', ''))
    disp  = _norm(entry.get('dispatch_name', ''))
    letter_out = _norm(entry.get('outgoing_letter', ''))

    found_row  = None
    candidates = []

    for row in range(ws.max_row, 2, -1):
        r_obj   = _norm(ws.cell(row, 2).value)
        r_disp  = _norm(ws.cell(row, 3).value)
        r_out   = _norm(ws.cell(row, 4).value)

        match_count = (
            (r_obj  == obj   and bool(obj))  +
            (r_disp == disp  and bool(disp)) +
            (r_out  == letter_out and bool(letter_out))
        )

        if match_count == 3:
            found_row = row
            break
        elif match_count >= 2:
            candidates.append({
                'row': row,
                'object': ws.cell(row, 2).value,
                'dispatch': ws.cell(row, 3).value,
                'outgoing': ws.cell(row, 4).value,
                'matches': match_count,
            })

    if found_row:
        _write_registry_row(ws, found_row, incoming_letter_num, incoming_num, letter_date)
        wb.save(REGISTRY_PATH)

    candidates.sort(key=lambda x: x['matches'], reverse=True)
    return found_row, candidates


def write_registry_row_manual(row: int, incoming_letter_num: str,
                               incoming_num: str, letter_date):
    """Записывает в указанную строку реестра."""
    wb = openpyxl.load_workbook(REGISTRY_PATH)
    ws = wb['Реестр']
    _write_registry_row(ws, row, incoming_letter_num, incoming_num, letter_date)
    wb.save(REGISTRY_PATH)


def _write_registry_row(ws, row: int, incoming_letter_num: str,
                        incoming_num: str, letter_date):
    ws.cell(row, 9).value  = incoming_letter_num
    ws.cell(row, 10).value = incoming_num
    if letter_date:
        if isinstance(letter_date, datetime):
            ws.cell(row, 11).value = letter_date
        else:
            dt = parse_date(str(letter_date))
            ws.cell(row, 11).value = dt if dt else str(letter_date)


# ── Регистрация таблиц уставок (старая форма, сводная) ───────────────────────

def write_to_summary_old(entry: dict, incoming_letter_num: str,
                         incoming_num: str) -> tuple:
    """
    Ищет строку в Регистрации (сводная) по номеру таблицы и записывает
    письмо об исполнении. Листы — по годам ('2025', '2024', ...).
    Возвращает (found_row: int|None, sheet_name: str).
    """
    if not HAS_OPENPYXL:
        raise RuntimeError("openpyxl не установлен")
    if not os.path.exists(SUMMARY_PATH):
        raise FileNotFoundError(f"Сводная не найдена:\n{SUMMARY_PATH}")

    table_num = entry.get('table_number', '')
    if not table_num:
        return None, ''

    # Определяем год из номера таблицы: "25-379" → "2025"
    m = re.match(r'(\d{2})-', table_num)
    year_sheet = ('20' + m.group(1)) if m else None

    wb = openpyxl.load_workbook(SUMMARY_PATH)
    service = {'служебный', 'шаблон', 'template'}
    year_sheets = [s for s in wb.sheetnames if s.lower() not in service]

    ws = None
    if year_sheet and year_sheet in wb.sheetnames:
        ws = wb[year_sheet]
    elif year_sheets:
        ws = wb[year_sheets[-1]]
    else:
        return None, ''

    # Находим столбцы по заголовкам строки 1
    num_col = None
    exec_col = None
    for col in range(1, ws.max_column + 1):
        h = _norm(ws.cell(1, col).value)
        if '№' in h and ('таблиц' in h or 'задан' in h):
            num_col = col
        elif 'исполнен' in h and 'письмо' in h:
            exec_col = col

    if not num_col or not exec_col:
        return None, ws.title

    # Ищем строку с нашим номером таблицы (снизу вверх)
    tnum_low = table_num.strip().lower()
    for row in range(ws.max_row, 1, -1):
        cell_val = _norm(ws.cell(row, num_col).value)
        if tnum_low in cell_val or cell_val == tnum_low:
            ws.cell(row, exec_col).value = f"{incoming_letter_num} вх-{incoming_num}"
            wb.save(SUMMARY_PATH)
            return row, ws.title

    return None, ws.title


# ── Файловые операции: поиск папок объектов ──────────────────────────────────

def _normalize_name(s: str) -> str:
    """Приводит к нижнему регистру, убирает ё→е, лишние пробелы."""
    return re.sub(r'\s+', ' ', s.lower().replace('ё', 'е').replace('й', 'й')).strip()


def match_object_to_short_name(object_name: str) -> str:
    """
    Сопоставляет официальное название объекта с коротким именем папки.
    Сначала обрезает 'ПС NNN кВ', затем использует SequenceMatcher.
    """
    if not object_name:
        return ''

    stripped = re.sub(r'^(ПС|ОРУ|ПП)\s+\d+\s*кВ\s+', '', object_name, flags=re.IGNORECASE).strip()

    # Точное совпадение после обрезки
    stripped_norm = _normalize_name(stripped)
    for name in OBJECT_SHORT_NAMES:
        if _normalize_name(name) == stripped_norm:
            return name

    # Fuzzy — сравниваем object_name и stripped со всеми short names
    best_name, best_score = '', 0.0
    for name in OBJECT_SHORT_NAMES:
        norm_name = _normalize_name(name)
        for candidate in [_normalize_name(stripped), _normalize_name(object_name)]:
            ratio = difflib.SequenceMatcher(None, candidate, norm_name).ratio()
            if ratio > best_score:
                best_score = ratio
                best_name = name

    return best_name if best_score >= 0.4 else ''


def find_object_exec_folder(short_name: str) -> str | None:
    """Возвращает путь к папке объекта в папке для исполнения."""
    if not short_name:
        return None
    direct = os.path.join(USTAVKI_EXEC_BASE, short_name)
    if os.path.isdir(direct):
        return direct
    # Поиск без учёта регистра
    try:
        for entry in os.scandir(USTAVKI_EXEC_BASE):
            if entry.is_dir() and _normalize_name(entry.name) == _normalize_name(short_name):
                return entry.path
    except OSError:
        pass
    return None


def find_current_and_archive_folders(object_folder: str) -> tuple:
    """Ищет подпапки 'Текущие' и 'Архив' в папке объекта."""
    current_dir = archive_dir = None
    if not object_folder or not os.path.isdir(object_folder):
        return None, None
    try:
        for entry in os.scandir(object_folder):
            if not entry.is_dir():
                continue
            n = entry.name.lower()
            if 'текущ' in n:
                current_dir = entry.path
            elif 'архив' in n:
                archive_dir = entry.path
    except OSError:
        pass
    return current_dir, archive_dir


def find_archive_candidates(new_filepath: str, current_dir: str, top_n: int = 5) -> list:
    """
    Ищет в current_dir файлы, похожие по названию на new_filepath.
    Обрезает последние 5 символов стема перед сравнением (шум).
    Возвращает list of (path, name, score).
    """
    if not current_dir or not os.path.isdir(current_dir):
        return []

    new_stem = os.path.splitext(os.path.basename(new_filepath))[0]
    new_trimmed = _normalize_name(new_stem[:-5] if len(new_stem) > 5 else new_stem)

    candidates = []
    try:
        for entry in os.scandir(current_dir):
            if not entry.is_file():
                continue
            stem = os.path.splitext(entry.name)[0]
            trimmed = _normalize_name(stem[:-5] if len(stem) > 5 else stem)
            score = difflib.SequenceMatcher(None, new_trimmed, trimmed).ratio()
            candidates.append((entry.path, entry.name, score))
    except OSError:
        pass

    candidates.sort(key=lambda x: x[2], reverse=True)
    return candidates[:top_n]


def move_table_files(entry: dict, archive_dir: str, current_dir: str) -> str:
    """
    Перекладывает архивный файл в Архив, новый файл в Текущие.
    Возвращает новый путь файла в Текущие.
    """
    archive_src = entry.get('archive_candidate', '')
    if archive_src and os.path.exists(archive_src):
        dest = os.path.join(archive_dir, os.path.basename(archive_src))
        shutil.move(archive_src, dest)

    src = entry.get('file_path', '')
    if src and os.path.exists(src):
        dest_new = os.path.join(current_dir, os.path.basename(src))
        shutil.copy2(src, dest_new)
        return dest_new
    return ''


# ── Visio: обновление карт уставок ───────────────────────────────────────────

def update_visio_map(visio_path: str, old_table_path: str,
                     new_table_path: str, new_table_number: str) -> tuple:
    """
    Открывает Visio, заменяет гиперссылку с old_table_path на new_table_path,
    сохраняет .vsdx и экспортирует PDF.
    Возвращает (success: bool, message: str).
    """
    try:
        import win32com.client as win32
    except ImportError:
        return False, "win32com не доступен"
    if not os.path.exists(visio_path):
        return False, f"Файл не найден: {visio_path}"

    visio = None
    try:
        visio = win32.Dispatch('Visio.Application')
        visio.Visible = False
        doc = visio.Documents.Open(os.path.abspath(visio_path))

        replaced = 0
        for page in doc.Pages:
            for shape in page.Shapes:
                for i in range(1, shape.Hyperlinks.Count + 1):
                    hl = shape.Hyperlinks.Item(i)
                    if old_table_path.lower() in hl.Address.lower():
                        hl.Address = new_table_path
                        if new_table_number:
                            hl.Description = new_table_number
                        replaced += 1

        doc.Save()

        # Экспорт в PDF
        stem = os.path.splitext(os.path.basename(visio_path))[0]
        pdf_path = os.path.join(MAPS_PDF_FOLDER, stem + '.pdf')
        os.makedirs(MAPS_PDF_FOLDER, exist_ok=True)
        doc.ExportAsFixedFormat(1, pdf_path, 0, 0)
        doc.Close()
        return True, f"Заменено ссылок: {replaced}, PDF: {pdf_path}"
    except Exception as exc:
        return False, str(exc)
    finally:
        try:
            if visio:
                visio.Quit()
        except Exception:
            pass


# ── ДЭБ: загрузка карт уставок ───────────────────────────────────────────────

def upload_to_deb_entry(dispatch_name: str, visio_path: str, pdf_path: str) -> tuple:
    """
    Загружает обновлённые файлы карты уставок в ДЭБ через Selenium + pyautogui.
    Возвращает (success: bool, message: str).
    """
    if not HAS_SELENIUM:
        return False, "selenium не установлен"
    if not HAS_PYAUTOGUI:
        return False, "pyautogui не установлен"

    try:
        opts = ChromeOptions()
        driver = webdriver.Chrome(options=opts)
        wait = WebDriverWait(driver, 30)

        # 1. Открываем каталог
        driver.get(DEB_MAPS_URL)
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'table tbody tr')))

        # 2. Ищем строку с нашим объектом по тексту
        # Текст в td: "Карта уставок ПС 220 кВ ..."
        rows = driver.find_elements(By.CSS_SELECTOR, 'table tbody tr')
        target_row = None
        dn_low = dispatch_name.lower()
        for row in rows:
            try:
                tds = row.find_elements(By.TAG_NAME, 'td')
                for td in tds:
                    if dn_low in td.text.lower():
                        target_row = row
                        break
            except Exception:
                pass
            if target_row:
                break

        if not target_row:
            driver.quit()
            return False, f"Объект не найден в ДЭБ: {dispatch_name}"

        # 3. Переходим по ссылке (кнопка "Перейти по ссылке")
        link_btn = target_row.find_element(By.CSS_SELECTOR, 'a[title="Перейти по ссылке"]')
        card_url = DEB_BASE_URL + link_btn.get_attribute('href').lstrip('.')
        driver.get(card_url)
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, '.edit-mode-button')))

        # 4. Нажимаем кнопку редактирования
        driver.find_element(By.CSS_SELECTOR, 'button.edit-mode-button').click()
        time.sleep(1.5)

        # 5. Обрабатываем файловые строки (visio и pdf)
        file_items = driver.find_elements(By.CSS_SELECTOR, '.list-group-item')
        for item in file_items:
            try:
                filename_el = item.find_element(By.CSS_SELECTOR, 'a.dz-filename')
                fname_text = filename_el.text.lower()
                is_visio = any(ext in fname_text for ext in ('.vsd', '.vsdx'))
                is_pdf   = '.pdf' in fname_text
                upload_path = visio_path if is_visio else (pdf_path if is_pdf else None)
                if not upload_path or not os.path.exists(upload_path):
                    continue

                change_btn = item.find_element(By.CSS_SELECTOR, 'a.change-file-button')
                change_btn.click()
                time.sleep(0.8)

                # Подтверждение (только для первого файла)
                try:
                    confirm_btn = wait.until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, 'button.btn.btn-success')),
                    )
                    confirm_btn.click()
                    time.sleep(0.8)
                except Exception:
                    pass

                # Выбор файла через pyautogui
                time.sleep(1.5)
                pyautogui.hotkey('ctrl', 'a')
                pyautogui.typewrite(upload_path, interval=0.02)
                pyautogui.press('enter')
                time.sleep(2.0)

            except Exception:
                continue

        # 6. Сохранить
        save_btn = driver.find_element(By.CSS_SELECTOR, 'button.confirm-button.btn-outline-success')
        save_btn.click()
        time.sleep(2.0)
        driver.quit()
        return True, f"Загружено для: {dispatch_name}"

    except Exception as exc:
        try:
            driver.quit()
        except Exception:
            pass
        return False, str(exc)


# ── UI: вспомогательный диалог "Строка не найдена" ───────────────────────────

class _RegistryNotFoundDialog(tk.Toplevel):
    """Диалог выбора строки реестра если автопоиск не нашёл совпадения."""

    def __init__(self, parent, entry_desc: str, candidates: list):
        super().__init__(parent)
        self.title("Строка не найдена в реестре")
        self.resizable(True, False)
        self.grab_set()
        self.result_row = None  # None = пропустить

        ttk.Label(self, text=f"Не найдено: {entry_desc}", wraplength=480,
                  font=('', 10, 'bold')).pack(padx=12, pady=(10, 4))

        if candidates:
            ttk.Label(self, text="Похожие строки (выберите или введите номер строки):").pack(padx=12)
            frame = ttk.Frame(self)
            frame.pack(fill='both', expand=True, padx=12, pady=4)
            cols = ('row', 'object', 'dispatch', 'outgoing')
            tv = ttk.Treeview(frame, columns=cols, show='headings', height=6)
            tv.heading('row', text='Строка'); tv.column('row', width=60)
            tv.heading('object', text='Объект'); tv.column('object', width=160)
            tv.heading('dispatch', text='Дисп. наим.'); tv.column('dispatch', width=200)
            tv.heading('outgoing', text='Исх. письмо'); tv.column('outgoing', width=120)
            for c in candidates:
                tv.insert('', 'end', values=(
                    c['row'], c.get('object',''), c.get('dispatch',''), c.get('outgoing','')
                ))
            sb = ttk.Scrollbar(frame, orient='vertical', command=tv.yview)
            tv.configure(yscrollcommand=sb.set)
            tv.pack(side='left', fill='both', expand=True)
            sb.pack(side='right', fill='y')
            tv.bind('<<TreeviewSelect>>', lambda e: self._tv_select(tv))
            self._tv = tv

        row_frame = ttk.Frame(self)
        row_frame.pack(padx=12, pady=4)
        ttk.Label(row_frame, text="Номер строки вручную:").pack(side='left')
        self._row_var = tk.StringVar()
        ttk.Entry(row_frame, textvariable=self._row_var, width=8).pack(side='left', padx=4)

        btn_frame = ttk.Frame(self)
        btn_frame.pack(pady=8)
        ttk.Button(btn_frame, text="Принять", command=self._on_accept).pack(side='left', padx=6)
        ttk.Button(btn_frame, text="Пропустить", command=self.destroy).pack(side='left', padx=6)

        self.transient(parent)
        self.wait_window()

    def _tv_select(self, tv):
        sel = tv.selection()
        if sel:
            val = tv.item(sel[0], 'values')
            self._row_var.set(val[0] if val else '')

    def _on_accept(self):
        try:
            self.result_row = int(self._row_var.get())
        except ValueError:
            messagebox.showwarning("Ошибка", "Введите целое число.", parent=self)
            return
        self.destroy()


# ── UI: главное окно ──────────────────────────────────────────────────────────

class RegistrationApp(_BASE_CLASS):
    """Главное окно: Регистрация писем + Таблицы уставок."""

    def __init__(self):
        super().__init__()
        self.in_data  = {}
        self.out_data = {}
        self.ustavki_entries: list = []   # list of dict
        self._in_preview_vars  = {}
        self._out_preview_vars = {}
        self._default_filename_in  = ''
        self._default_filename_out = ''
        self._ustavki_files_listbox = None
        self._tree = None
        self._archive_tree = None
        self._step_log_widgets: dict = {}  # step_name → tk.Text

        self.title("Регистрация корреспонденции v2")
        self.resizable(True, True)
        self._build_ui()
        self._center_window()

    # ─── Построение UI ────────────────────────────────────────────────────

    def _build_ui(self):
        root_frame = ttk.Frame(self, padding=8)
        root_frame.grid(row=0, column=0, sticky='nsew')
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        root_frame.columnconfigure(0, weight=1)
        root_frame.rowconfigure(0, weight=1)

        self._top_nb = ttk.Notebook(root_frame)
        self._top_nb.grid(row=0, column=0, sticky='nsew')
        self._top_nb.bind('<<NotebookTabChanged>>', self._on_top_tab_changed)

        # ── Вкладка 1: Регистрация писем ──
        reg_tab = ttk.Frame(self._top_nb, padding=4)
        self._top_nb.add(reg_tab, text="  Регистрация писем  ")
        reg_tab.columnconfigure(0, weight=1)
        reg_tab.rowconfigure(0, weight=1)

        self._reg_nb = ttk.Notebook(reg_tab)
        self._reg_nb.grid(row=0, column=0, sticky='nsew')

        in_tab  = ttk.Frame(self._reg_nb, padding=8)
        out_tab = ttk.Frame(self._reg_nb, padding=8)
        self._reg_nb.add(in_tab,  text="  Входящие  ")
        self._reg_nb.add(out_tab, text="  Исходящие  ")
        self._build_incoming_tab(in_tab)
        self._build_outgoing_tab(out_tab)

        # Статус и кнопки для вкладки регистрации
        self._reparse_status = tk.StringVar(value="")
        ttk.Label(reg_tab, textvariable=self._reparse_status,
                  foreground='gray').grid(row=1, column=0, pady=(2, 0))
        btn_frame = ttk.Frame(reg_tab)
        btn_frame.grid(row=2, column=0, pady=4)
        ttk.Button(btn_frame, text="Зарегистрировать в журнал",
                   command=self._on_register).pack(side='left', padx=6)
        self._reparse_btn = ttk.Button(btn_frame, text="Запустить парсинг LanDocs",
                                       command=self._start_reparse)
        self._reparse_btn.pack(side='left', padx=6)
        ttk.Button(btn_frame, text="Закрыть",
                   command=self.destroy).pack(side='left', padx=6)

        # ── Вкладка 2: Таблицы уставок ──
        ust_tab = ttk.Frame(self._top_nb, padding=4)
        self._top_nb.add(ust_tab, text="  Таблицы уставок  ")
        ust_tab.columnconfigure(0, weight=1)
        ust_tab.rowconfigure(1, weight=1)
        self._build_ustavki_tab(ust_tab)

    def _build_incoming_tab(self, frame):
        frame.columnconfigure(1, weight=1)

        # Блок импорта из журнала
        imp = ttk.LabelFrame(frame, text="Импорт из журнала", padding=8)
        imp.grid(row=0, column=0, sticky='ew', pady=(0, 6))
        imp.columnconfigure(3, weight=1)
        ttk.Label(imp, text="№ вх:").grid(row=0, column=0, sticky='e', padx=(0,4))
        self._in_import_num = tk.StringVar()
        ttk.Entry(imp, textvariable=self._in_import_num, width=18).grid(row=0, column=1)
        ttk.Label(imp, text="  Год:").grid(row=0, column=2, sticky='e', padx=(8,4))
        self._in_import_year = tk.StringVar(value=str(datetime.now().year))
        ttk.Entry(imp, textvariable=self._in_import_year, width=6).grid(row=0, column=3, sticky='w')
        ttk.Button(imp, text="Найти в журнале",
                   command=self._on_import_journal).grid(row=0, column=4, padx=(8,0))
        self._import_status = tk.StringVar(value="")
        ttk.Label(imp, textvariable=self._import_status,
                  foreground='navy').grid(row=1, column=0, columnspan=5, pady=(4,0), sticky='w')

        # Данные из LanDocs
        info = ttk.LabelFrame(frame, text="Данные из LanDocs", padding=8)
        info.grid(row=1, column=0, sticky='ew', pady=(0, 8))
        info.columnconfigure(1, weight=1)
        fields = [
            ("Дата:",             'date'),
            ("№ вх:",             'incoming_num'),
            ("№ письма:",         'letter_num'),
            ("Тема письма:",      'subject'),
            ("Подписант:",        'signatory'),
            ("Корреспондент:",    'correspondent'),
            ("Связанное письмо:", 'related'),
        ]
        for i, (label, key) in enumerate(fields):
            ttk.Label(info, text=label, anchor='e').grid(
                row=i, column=0, sticky='e', padx=(0, 6), pady=2)
            var = tk.StringVar()
            self._in_preview_vars[key] = var
            ttk.Label(info, textvariable=var, anchor='w', wraplength=420).grid(
                row=i, column=1, sticky='w', pady=2)

        # Данные для регистрации
        inp = ttk.LabelFrame(frame, text="Данные для регистрации", padding=8)
        inp.grid(row=2, column=0, sticky='ew')
        inp.columnconfigure(1, weight=1)
        ttk.Label(inp, text="Автор письма:").grid(row=0, column=0, sticky='e', padx=(0,6), pady=4)
        self.in_author_var = tk.StringVar(value="-")
        ttk.Entry(inp, textvariable=self.in_author_var, width=48).grid(row=0, column=1, sticky='ew', pady=4)
        ttk.Label(inp, text="Ключевые слова:").grid(row=1, column=0, sticky='e', padx=(0,6), pady=4)
        self.in_keywords_var = tk.StringVar()
        ttk.Entry(inp, textvariable=self.in_keywords_var, width=48).grid(row=1, column=1, sticky='ew', pady=4)
        ttk.Label(inp, text="Название файла:").grid(row=2, column=0, sticky='e', padx=(0,6), pady=4)
        self.in_filename_var = tk.StringVar()
        ttk.Entry(inp, textvariable=self.in_filename_var, width=48).grid(row=2, column=1, sticky='ew', pady=4)
        ttk.Label(inp, text="Папка сохранения:").grid(row=3, column=0, sticky='e', padx=(0,6), pady=4)
        row3 = ttk.Frame(inp); row3.grid(row=3, column=1, sticky='ew', pady=4)
        row3.columnconfigure(0, weight=1)
        self.in_save_path_var = tk.StringVar()
        ttk.Entry(row3, textvariable=self.in_save_path_var, state='readonly', width=38).grid(row=0, column=0, sticky='ew')
        ttk.Button(row3, text="Выбрать папку…",
                   command=self._choose_save_folder_in).grid(row=0, column=1, padx=(6,0))
        ttk.Label(inp, text="№ папки:").grid(row=4, column=0, sticky='e', padx=(0,6), pady=4)
        self.in_folder_num_var = tk.StringVar()
        ttk.Label(inp, textvariable=self.in_folder_num_var,
                  anchor='w', foreground='navy').grid(row=4, column=1, sticky='w', pady=4)

    def _build_outgoing_tab(self, frame):
        frame.columnconfigure(1, weight=1)
        info = ttk.LabelFrame(frame, text="Данные из LanDocs", padding=8)
        info.grid(row=0, column=0, sticky='ew', pady=(0, 8))
        info.columnconfigure(1, weight=1)
        fields = [
            ("Дата:",             'date'),
            ("№ письма:",         'letter_num'),
            ("Тема письма:",      'subject'),
            ("Исполнитель:",      'executor'),
            ("Получатели (ФИО):", 'recipient_names'),
            ("Получатели (орг):", 'recipient_companies'),
            ("Связанное письмо:", 'related'),
        ]
        for i, (label, key) in enumerate(fields):
            ttk.Label(info, text=label, anchor='e').grid(
                row=i, column=0, sticky='e', padx=(0, 6), pady=2)
            var = tk.StringVar()
            self._out_preview_vars[key] = var
            ttk.Label(info, textvariable=var, anchor='w', wraplength=420).grid(
                row=i, column=1, sticky='w', pady=2)
        inp = ttk.LabelFrame(frame, text="Данные для регистрации", padding=8)
        inp.grid(row=1, column=0, sticky='ew')
        inp.columnconfigure(1, weight=1)
        ttk.Label(inp, text="Ключевые слова:").grid(row=0, column=0, sticky='e', padx=(0,6), pady=4)
        self.out_keywords_var = tk.StringVar()
        ttk.Entry(inp, textvariable=self.out_keywords_var, width=48).grid(row=0, column=1, sticky='ew', pady=4)
        ttk.Label(inp, text="Контроль:").grid(row=1, column=0, sticky='e', padx=(0,6), pady=4)
        self.out_control_var = tk.StringVar()
        ttk.Entry(inp, textvariable=self.out_control_var, width=48).grid(row=1, column=1, sticky='ew', pady=4)
        ttk.Label(inp, text="Название файла:").grid(row=2, column=0, sticky='e', padx=(0,6), pady=4)
        self.out_filename_var = tk.StringVar()
        ttk.Entry(inp, textvariable=self.out_filename_var, width=48).grid(row=2, column=1, sticky='ew', pady=4)
        ttk.Label(inp, text="Папка сохранения:").grid(row=3, column=0, sticky='e', padx=(0,6), pady=4)
        row3 = ttk.Frame(inp); row3.grid(row=3, column=1, sticky='ew', pady=4)
        row3.columnconfigure(0, weight=1)
        self.out_save_path_var = tk.StringVar()
        ttk.Entry(row3, textvariable=self.out_save_path_var, state='readonly', width=38).grid(row=0, column=0, sticky='ew')
        ttk.Button(row3, text="Выбрать папку…",
                   command=self._choose_save_folder_out).grid(row=0, column=1, padx=(6,0))
        ttk.Label(inp, text="№ папки:").grid(row=4, column=0, sticky='e', padx=(0,6), pady=4)
        self.out_folder_num_var = tk.StringVar()
        ttk.Label(inp, textvariable=self.out_folder_num_var,
                  anchor='w', foreground='navy').grid(row=4, column=1, sticky='w', pady=4)


    def _build_ustavki_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)

        # Верхняя полоска: информация о письме
        letter_frame = ttk.LabelFrame(parent, text="Письмо о выставлении", padding=6)
        letter_frame.grid(row=0, column=0, sticky='ew', padx=4, pady=(4,0))
        letter_frame.columnconfigure(3, weight=1)
        ttk.Label(letter_frame, text="№ вх:").grid(row=0, column=0, sticky='e', padx=(0,3))
        self._lu_vx    = tk.StringVar(); ttk.Label(letter_frame, textvariable=self._lu_vx,
            foreground='navy', width=12).grid(row=0, column=1, sticky='w')
        ttk.Label(letter_frame, text="№ письма:").grid(row=0, column=2, sticky='e', padx=(10,3))
        self._lu_letter = tk.StringVar(); ttk.Label(letter_frame, textvariable=self._lu_letter,
            foreground='navy', width=25).grid(row=0, column=3, sticky='w')
        ttk.Label(letter_frame, text="Дата:").grid(row=0, column=4, sticky='e', padx=(10,3))
        self._lu_date  = tk.StringVar(); ttk.Label(letter_frame, textvariable=self._lu_date,
            foreground='navy', width=12).grid(row=0, column=5, sticky='w')
        ttk.Label(letter_frame, text="Тема:").grid(row=1, column=0, sticky='e', padx=(0,3))
        self._lu_subj  = tk.StringVar(); ttk.Label(letter_frame, textvariable=self._lu_subj,
            foreground='gray', wraplength=700, anchor='w').grid(
            row=1, column=1, columnspan=5, sticky='w', pady=(2,0))

        # Ступенчатые вкладки шагов
        self._step_nb = ttk.Notebook(parent)
        self._step_nb.grid(row=1, column=0, sticky='nsew', padx=4, pady=4)

        step0 = ttk.Frame(self._step_nb, padding=8); self._step_nb.add(step0, text=" 0 Файлы ")
        step1 = ttk.Frame(self._step_nb, padding=8); self._step_nb.add(step1, text=" 1 Данные ")
        step2 = ttk.Frame(self._step_nb, padding=8); self._step_nb.add(step2, text=" 2 Запись в таблицы ")
        step3 = ttk.Frame(self._step_nb, padding=8); self._step_nb.add(step3, text=" 3 Реестры ")
        step4 = ttk.Frame(self._step_nb, padding=8); self._step_nb.add(step4, text=" 4 Раскладка ")
        step5 = ttk.Frame(self._step_nb, padding=8); self._step_nb.add(step5, text=" 5 Изменения ")
        step6 = ttk.Frame(self._step_nb, padding=8); self._step_nb.add(step6, text=" 6 Карты ")
        step7 = ttk.Frame(self._step_nb, padding=8); self._step_nb.add(step7, text=" 7 ДЭБ ")

        self._build_step0(step0)
        self._build_step1(step1)
        self._build_step_log(step2, " 2 Запись в таблицы",
            "Запишет «Уставки выставлены: ПИСЬМО вх-N от ДАТА» в каждый файл .docx",
            "Записать в таблицы", self._write_issued_all)
        self._build_step_log(step3, " 3 Реестры",
            "Новая форма → Реестр таблиц уставок.xlsx  |  Старая форма → Регистрация таблиц уставок.xlsx",
            "Записать в реестры", self._write_registries_all)
        self._build_step4(step4)
        self._build_step_log(step5, " 5 Изменения",
            "Ищет строки синего цвета в таблицах уставок и формирует отчёт Word",
            "Создать отчёт изменений", self._create_changes_report)
        self._build_step_log(step6, " 6 Карты",
            "Обновляет гиперссылки в картах уставок Visio и экспортирует PDF.\nТребует: Microsoft Visio",
            "Обновить карты уставок", self._update_visio_maps_all)
        self._build_step_log(step7, " 7 ДЭБ",
            "Загружает обновлённые карты в ДЭБ через браузер.\nТребует: selenium, pyautogui, ChromeDriver",
            "Загрузить в ДЭБ", self._upload_deb_all)

    def _build_step0(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(2, weight=1)

        # Зона drag-and-drop
        dnd_frame = ttk.LabelFrame(parent, text="Перетащите таблицы уставок (.docx)", padding=8)
        dnd_frame.grid(row=0, column=0, sticky='ew', pady=(0,6))
        dnd_label = ttk.Label(dnd_frame,
            text="Перетащите файлы .docx/.doc сюда\nили нажмите кнопку ниже",
            justify='center', foreground='gray')
        dnd_label.pack(pady=12)

        if HAS_DND:
            dnd_frame.drop_target_register(DND_FILES)
            dnd_frame.dnd_bind('<<Drop>>', self._on_dnd_drop)
            dnd_label.configure(text="Перетащите файлы .docx/.doc сюда  (или кнопку ниже)")

        # Кнопки
        btn_frame = ttk.Frame(parent)
        btn_frame.grid(row=1, column=0, sticky='ew', pady=(0,4))
        ttk.Button(btn_frame, text="Добавить файлы...",
                   command=self._add_ustavki_files).pack(side='left', padx=4)
        ttk.Button(btn_frame, text="Удалить выбранные",
                   command=self._remove_selected_files).pack(side='left', padx=4)
        ttk.Button(btn_frame, text="→ Парсировать",
                   command=self._go_parse).pack(side='right', padx=4)

        # Список файлов
        list_frame = ttk.LabelFrame(parent, text="Добавленные файлы", padding=4)
        list_frame.grid(row=2, column=0, sticky='nsew')
        list_frame.columnconfigure(0, weight=1)
        list_frame.rowconfigure(0, weight=1)
        sb = ttk.Scrollbar(list_frame, orient='vertical')
        self._ustavki_files_listbox = tk.Listbox(list_frame, yscrollcommand=sb.set,
                                                   selectmode='extended', height=10)
        sb.configure(command=self._ustavki_files_listbox.yview)
        self._ustavki_files_listbox.grid(row=0, column=0, sticky='nsew')
        sb.grid(row=0, column=1, sticky='ns')

    def _build_step1(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)

        # Treeview
        tree_frame = ttk.Frame(parent)
        tree_frame.grid(row=0, column=0, sticky='nsew')
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)

        cols = ('file','form','object','dispatch','table_num',
                'outgoing_letter','letter_num','letter_date','status')
        self._tree = ttk.Treeview(tree_frame, columns=cols, show='headings', height=14)
        widths = [160, 55, 160, 200, 90, 140, 140, 100, 90]
        heads  = ['Файл','Форма','Объект','Дисп. наим.',
                  '№ таблицы','Исх. письмо','Вх. письмо','Дата вх.','Статус']
        for col, head, w in zip(cols, heads, widths):
            self._tree.heading(col, text=head)
            self._tree.column(col, width=w, minwidth=40)
        vsb = ttk.Scrollbar(tree_frame, orient='vertical',   command=self._tree.yview)
        hsb = ttk.Scrollbar(tree_frame, orient='horizontal', command=self._tree.xview)
        self._tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self._tree.grid(row=0, column=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        hsb.grid(row=1, column=0, sticky='ew')
        self._tree.bind('<Double-1>', self._on_tree_double_click)

        # Кнопки
        btn_frame = ttk.Frame(parent)
        btn_frame.grid(row=1, column=0, sticky='ew', pady=(6,0))
        ttk.Button(btn_frame, text="Парсировать все",
                   command=self._parse_ustavki_all).pack(side='left', padx=4)
        ttk.Button(btn_frame, text="Применить письмо из журнала",
                   command=self._apply_letter_to_all).pack(side='left', padx=4)
        ttk.Button(btn_frame, text="Очистить список",
                   command=self._clear_entries).pack(side='right', padx=4)

    def _build_step_log(self, parent, title: str, description: str,
                        btn_text: str, btn_command):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)
        ttk.Label(parent, text=description, wraplength=700,
                  foreground='gray').grid(row=0, column=0, sticky='w', pady=(0,6))
        txt = tk.Text(parent, height=12, wrap='word', state='disabled',
                      font=('Consolas', 9))
        sb = ttk.Scrollbar(parent, orient='vertical', command=txt.yview)
        txt.configure(yscrollcommand=sb.set)
        txt.grid(row=1, column=0, sticky='nsew')
        sb.grid(row=1, column=1, sticky='ns')
        ttk.Button(parent, text=btn_text, command=btn_command).grid(
            row=2, column=0, pady=(6,0), sticky='w')
        self._step_log_widgets[title.strip()] = txt

    def _build_step4(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)
        ttk.Label(parent,
            text="Ищет похожие файлы в папках «Текущие», предлагает кандидатов для архивации.",
            wraplength=700, foreground='gray').grid(row=0, column=0, sticky='w', pady=(0,4))

        arc_frame = ttk.Frame(parent)
        arc_frame.grid(row=1, column=0, sticky='nsew')
        arc_frame.columnconfigure(0, weight=1)
        arc_frame.rowconfigure(0, weight=1)

        arc_cols = ('file','short_name','archive_candidate','score')
        self._archive_tree = ttk.Treeview(arc_frame, columns=arc_cols,
                                           show='headings', height=12)
        self._archive_tree.heading('file',             text='Новый файл')
        self._archive_tree.heading('short_name',       text='Объект (папка)')
        self._archive_tree.heading('archive_candidate',text='Кандидат на архив')
        self._archive_tree.heading('score',            text='Схожесть')
        self._archive_tree.column('file',             width=180)
        self._archive_tree.column('short_name',       width=120)
        self._archive_tree.column('archive_candidate',width=220)
        self._archive_tree.column('score',            width=70)
        vsb = ttk.Scrollbar(arc_frame, orient='vertical', command=self._archive_tree.yview)
        hsb = ttk.Scrollbar(arc_frame, orient='horizontal', command=self._archive_tree.xview)
        self._archive_tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self._archive_tree.grid(row=0, column=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        hsb.grid(row=1, column=0, sticky='ew')
        self._archive_tree.bind('<Double-1>', self._on_archive_tree_dclick)

        btn_frame = ttk.Frame(parent)
        btn_frame.grid(row=2, column=0, pady=(6,0), sticky='w')
        ttk.Button(btn_frame, text="Найти кандидатов",
                   command=self._find_archive_candidates_all).pack(side='left', padx=4)
        ttk.Button(btn_frame, text="Разложить файлы",
                   command=self._move_files_all).pack(side='left', padx=4)


    # ─── Обновление инфо-полосы письма ────────────────────────────────────

    def _refresh_letter_info(self):
        d = self.in_data
        self._lu_vx.set(d.get('incoming_num', ''))
        self._lu_letter.set(d.get('letter_num', ''))
        self._lu_date.set(d.get('date', ''))
        self._lu_subj.set(d.get('subject', ''))

    def _on_top_tab_changed(self, _event=None):
        if self._top_nb.index(self._top_nb.select()) == 1:
            self._refresh_letter_info()

    # ─── Импорт из журнала ────────────────────────────────────────────────

    def _on_import_journal(self):
        num  = self._in_import_num.get().strip()
        year = self._in_import_year.get().strip()
        if not num:
            self._import_status.set("Введите № вх")
            return
        try:
            result = lookup_incoming_journal(num, year)
        except Exception as exc:
            self._import_status.set(f"Ошибка: {exc}")
            return
        if result is None:
            self._import_status.set(f"Не найдено: {num} за {year} год")
            return
        self.in_data = {
            'incoming_num': result.get('incoming_num', num),
            'date':         result.get('date', ''),
            'letter_num':   result.get('letter_num', ''),
            'subject':      result.get('subject', ''),
            'signatory':    '',
            'correspondent': result.get('signed_by', ''),
            'related':      result.get('related', ''),
            'file_link':    result.get('hyperlink', ''),
        }
        self._apply_incoming_data()
        self.in_author_var.set(result.get('author', '-'))
        self.in_keywords_var.set(result.get('keywords', ''))
        self._import_status.set(f"Импортировано: {num}")
        self._refresh_letter_info()

    # ─── Применение данных LanDocs ────────────────────────────────────────

    def _apply_incoming_data(self):
        d = self.in_data
        for key, var in self._in_preview_vars.items():
            var.set(d.get(key, ''))
        self._default_filename_in = build_default_filename_in(
            d.get('date',''), d.get('incoming_num',''), d.get('letter_num',''))
        self.in_filename_var.set(self._default_filename_in)
        self.in_save_path_var.set('')
        self.in_folder_num_var.set('')

    def _apply_outgoing_data(self):
        d = self.out_data
        for key, var in self._out_preview_vars.items():
            var.set(d.get(key, ''))
        self._default_filename_out = build_default_filename_out(
            d.get('date',''), d.get('letter_num',''))
        self.out_filename_var.set(self._default_filename_out)
        self.out_keywords_var.set(d.get('subject',''))
        self.out_save_path_var.set('')
        self.out_folder_num_var.set('')

    # ─── Парсинг LanDocs ──────────────────────────────────────────────────

    def _active_reg_tab(self) -> str:
        return 'out' if self._reg_nb.index(self._reg_nb.select()) == 1 else 'in'

    def _start_reparse(self):
        self._reparse_btn.config(state='disabled')
        self.iconify()
        self._reparse_countdown(3)

    def _reparse_countdown(self, n: int):
        if n > 0:
            self._reparse_status.set(f"Переключитесь в LanDocs… парсинг через {n} сек.")
            self.after(1000, self._reparse_countdown, n - 1)
        else:
            self._reparse_status.set("Выполняется парсинг…")
            self.after(50, self._do_reparse)

    def _do_reparse(self):
        tab = self._active_reg_tab()
        try:
            if tab == 'in':
                self.in_data = extract_landocs_data_in()
                self._apply_incoming_data()
            else:
                self.out_data = extract_landocs_data_out()
                self._apply_outgoing_data()
            self._reparse_status.set("Парсинг завершён успешно.")
        except Exception as exc:
            self._reparse_status.set(f"Ошибка: {exc}")
        finally:
            self._reparse_btn.config(state='normal')
            self.deiconify()
            self.lift()

    # ─── Выбор папки / файла ──────────────────────────────────────────────

    def _choose_save_folder_in(self):
        self._choose_save_folder(
            file_link=self.in_data.get('file_link',''),
            filename_var=self.in_filename_var,
            default_filename=self._default_filename_in,
            save_path_var=self.in_save_path_var,
            folder_num_var=self.in_folder_num_var,
        )

    def _choose_save_folder_out(self):
        self._choose_save_folder(
            file_link=self.out_data.get('file_link',''),
            filename_var=self.out_filename_var,
            default_filename=self._default_filename_out,
            save_path_var=self.out_save_path_var,
            folder_num_var=self.out_folder_num_var,
        )

    def _choose_save_folder(self, file_link, filename_var, default_filename,
                            save_path_var, folder_num_var):
        initial = DEFAULT_SAVE_FOLDER if os.path.isdir(DEFAULT_SAVE_FOLDER) else os.path.expanduser("~")
        ext = os.path.splitext(file_link)[1].lower() if file_link else '.pdf'
        if not ext: ext = '.pdf'
        filename_base = filename_var.get() or default_filename
        selected = filedialog.asksaveasfilename(
            title="Выберите папку и имя файла для сохранения письма",
            initialdir=initial,
            initialfile=filename_base + ext,
            defaultextension=ext,
            filetypes=[("PDF файлы", "*.pdf"), ("Все файлы", "*.*")],
        )
        if not selected:
            return
        selected = selected.replace('/', '\\')
        save_path_var.set(selected)
        base_name = os.path.splitext(os.path.basename(selected))[0]
        if base_name:
            filename_var.set(base_name)
        folder_num_var.set(calc_folder_num(os.path.dirname(selected)))

    # ─── Регистрация в журнал ─────────────────────────────────────────────

    def _on_register(self):
        if self._active_reg_tab() == 'in':
            self._on_register_in()
        else:
            self._on_register_out()

    def _on_register_in(self):
        if not self.in_save_path_var.get():
            messagebox.showwarning("Внимание", "Выберите папку и имя файла.", parent=self)
            return
        try:
            d = self.in_data
            date_str      = d.get('date','')
            signatory     = d.get('signatory','')
            correspondent = d.get('correspondent','')
            signed_by = f"{signatory}\n{correspondent}" if correspondent else signatory
            hyperlink_path = self.in_save_path_var.get()
            if hyperlink_path:
                src = find_latest_in_viewdir()
                if src:
                    shutil.copy2(src, hyperlink_path)
                else:
                    raise FileNotFoundError(
                        "Файл письма не найден в ViewDir.\n"
                        r"Убедитесь, что письмо открыто в LanDocs (%LOCALAPPDATA%\Temp\ViewDir)")
            write_to_excel_in({
                'date':           fmt_date_ymd(date_str),
                'incoming_num':   d.get('incoming_num',''),
                'letter_num':     d.get('letter_num',''),
                'subject':        d.get('subject',''),
                'author':         self.in_author_var.get(),
                'signed_by':      signed_by,
                'folder_num':     self.in_folder_num_var.get(),
                'who_registered': getpass.getuser(),
                'keywords':       self.in_keywords_var.get(),
                'related':        d.get('related',''),
                'hyperlink_path': hyperlink_path,
            })
            messagebox.showinfo("Готово", "Запись добавлена в журнал!", parent=self)
            self._refresh_letter_info()
        except Exception as exc:
            messagebox.showerror("Ошибка", str(exc), parent=self)

    def _on_register_out(self):
        if not self.out_save_path_var.get():
            messagebox.showwarning("Внимание", "Выберите папку и имя файла.", parent=self)
            return
        try:
            d = self.out_data
            hyperlink_path = self.out_save_path_var.get()
            if hyperlink_path:
                src = find_latest_in_viewdir()
                if src:
                    shutil.copy2(src, hyperlink_path)
                else:
                    raise FileNotFoundError(
                        "Файл письма не найден в ViewDir.")
            write_to_excel_out({
                'date':           fmt_date_dmy(d.get('date','')),
                'letter_num':     d.get('letter_num',''),
                'subject':        d.get('subject',''),
                'recipient':      build_recipient_string(d.get('recipient_names',''),
                                                         d.get('recipient_companies','')),
                'executor':       d.get('executor',''),
                'keywords':       self.out_keywords_var.get(),
                'related':        d.get('related',''),
                'control':        self.out_control_var.get(),
                'hyperlink_path': hyperlink_path,
            })
            messagebox.showinfo("Готово", "Запись добавлена в журнал!", parent=self)
        except Exception as exc:
            messagebox.showerror("Ошибка", str(exc), parent=self)


    # ─── Шаг 0: файлы ─────────────────────────────────────────────────────

    def _on_dnd_drop(self, event):
        raw = event.data
        # tkinterdnd2 возвращает пути в фигурных скобках если есть пробелы
        paths = []
        for part in re.findall(r'\{([^}]+)\}|(\S+)', raw):
            p = part[0] or part[1]
            if p:
                paths.append(p)
        for p in paths:
            if p.lower().endswith(('.docx', '.doc')):
                self._add_entry(p)

    def _add_ustavki_files(self):
        initial = USTAVKI_EXEC_BASE if os.path.isdir(USTAVKI_EXEC_BASE) else os.path.expanduser('~')
        files = filedialog.askopenfilenames(
            title="Выберите таблицы уставок",
            initialdir=initial,
            filetypes=[("Word файлы", "*.docx *.doc"), ("Все файлы", "*.*")],
        )
        for f in files:
            self._add_entry(f)

    def _add_entry(self, path: str):
        # Проверить дубликат
        for e in self.ustavki_entries:
            if e['file_path'] == path:
                return
        self.ustavki_entries.append({
            'file_path':        path,
            'form_type':        '',
            'object_name':      '',
            'dispatch_name':    '',
            'table_number':     '',
            'outgoing_letter':  '',
            'letter_num':       '',
            'letter_date':      '',
            'status':           'ожидание',
            'registry_row':     0,
            'archive_candidate': '',
            'current_path':     '',
        })
        if self._ustavki_files_listbox:
            self._ustavki_files_listbox.insert('end', os.path.basename(path))

    def _remove_selected_files(self):
        lb = self._ustavki_files_listbox
        if not lb:
            return
        sel = list(lb.curselection())
        for idx in reversed(sel):
            lb.delete(idx)
            if idx < len(self.ustavki_entries):
                self.ustavki_entries.pop(idx)
        self._refresh_tree()

    def _go_parse(self):
        self._step_nb.select(1)
        self._parse_ustavki_all()

    # ─── Шаг 1: парсинг ───────────────────────────────────────────────────

    def _parse_ustavki_all(self):
        if not HAS_DOCX:
            messagebox.showerror("Ошибка", "python-docx не установлен.", parent=self)
            return
        for entry in self.ustavki_entries:
            try:
                parsed = parse_ustavki_table(entry['file_path'])
                entry.update({
                    'form_type':       parsed['form_type'],
                    'object_name':     parsed['object_name'],
                    'dispatch_name':   parsed['dispatch_name'],
                    'table_number':    parsed['table_number'],
                    'outgoing_letter': parsed['outgoing_letter'],
                    'status':          'спарсено',
                })
            except Exception as exc:
                entry['status'] = f'ошибка: {exc}'
        self._apply_letter_to_all()
        self._refresh_tree()

    def _apply_letter_to_all(self):
        lnum  = self.in_data.get('letter_num',  '')
        ldate = self.in_data.get('date',         '')
        lvx   = self.in_data.get('incoming_num', '')
        for entry in self.ustavki_entries:
            entry['letter_num']  = lnum
            entry['letter_date'] = ldate
            entry['letter_vx']   = lvx
        self._refresh_tree()
        self._refresh_letter_info()

    def _refresh_tree(self):
        if not self._tree:
            return
        self._tree.delete(*self._tree.get_children())
        for entry in self.ustavki_entries:
            self._tree.insert('', 'end', values=(
                os.path.basename(entry.get('file_path','')),
                entry.get('form_type',''),
                entry.get('object_name',''),
                entry.get('dispatch_name',''),
                entry.get('table_number',''),
                entry.get('outgoing_letter',''),
                entry.get('letter_num',''),
                entry.get('letter_date',''),
                entry.get('status',''),
            ))

    def _clear_entries(self):
        self.ustavki_entries.clear()
        if self._ustavki_files_listbox:
            self._ustavki_files_listbox.delete(0, 'end')
        self._refresh_tree()

    # ─── Inline-редактор ячеек Treeview ───────────────────────────────────

    # Карта: индекс столбца → ключ в entry dict
    _TREE_COL_KEYS = [
        None,            # file (readonly)
        'form_type',
        'object_name',
        'dispatch_name',
        'table_number',
        'outgoing_letter',
        'letter_num',
        'letter_date',
        'status',
    ]

    def _on_tree_double_click(self, event):
        tree = self._tree
        region = tree.identify_region(event.x, event.y)
        if region != 'cell':
            return
        col_id  = tree.identify_column(event.x)   # '#1', '#2', ...
        row_id  = tree.identify_row(event.y)
        if not row_id:
            return
        col_idx = int(col_id.lstrip('#')) - 1  # 0-based
        if col_idx == 0:  # file — не редактируем
            return
        # Находим запись
        all_ids = tree.get_children()
        try:
            entry_idx = list(all_ids).index(row_id)
        except ValueError:
            return
        if entry_idx >= len(self.ustavki_entries):
            return
        entry = self.ustavki_entries[entry_idx]
        key = self._TREE_COL_KEYS[col_idx] if col_idx < len(self._TREE_COL_KEYS) else None
        if not key:
            return

        # Координаты ячейки
        x, y, w, h = tree.bbox(row_id, col_id)
        if not w:
            return

        var = tk.StringVar(value=entry.get(key, ''))
        ent = ttk.Entry(tree, textvariable=var)
        ent.place(x=x, y=y, width=w, height=h)
        ent.focus_set()
        ent.select_range(0, 'end')

        def _commit(_event=None):
            entry[key] = var.get()
            # Обновляем значение в treeview
            vals = list(tree.item(row_id, 'values'))
            vals[col_idx] = var.get()
            tree.item(row_id, values=vals)
            ent.destroy()

        ent.bind('<Return>', _commit)
        ent.bind('<Escape>', lambda _: ent.destroy())
        ent.bind('<FocusOut>', _commit)

    # ─── Шаг 2: запись в таблицы ──────────────────────────────────────────

    def _log(self, step_key: str, text: str):
        widget = self._step_log_widgets.get(step_key)
        if widget:
            widget.configure(state='normal')
            widget.insert('end', text + '\n')
            widget.see('end')
            widget.configure(state='disabled')

    def _write_issued_all(self):
        log = '2 Запись в таблицы'
        if not HAS_DOCX:
            self._log(log, "ОШИБКА: python-docx не установлен")
            return
        vx    = self.in_data.get('incoming_num', '')
        lnum  = self.in_data.get('letter_num', '')
        ldate = self.in_data.get('date', '')
        hpath = self.in_data.get('file_link', '')

        for entry in self.ustavki_entries:
            fpath = entry.get('file_path', '')
            fname = os.path.basename(fpath)
            try:
                ok = write_issued_to_doc(fpath, lnum, vx, ldate, hpath)
                if ok:
                    entry['status'] = 'выставлено'
                    self._log(log, f"OK  {fname}")
                else:
                    entry['status'] = 'нет поля'
                    self._log(log, f"НЕТ ПОЛЯ 'Уставки выставлены'  {fname}")
            except Exception as exc:
                entry['status'] = f'ошибка'
                self._log(log, f"ОШИБКА  {fname}: {exc}")
        self._refresh_tree()

    # ─── Шаг 3: реестры ───────────────────────────────────────────────────

    def _write_registries_all(self):
        log = '3 Реестры'
        vx    = self.in_data.get('incoming_num', '')
        lnum  = self.in_data.get('letter_num', '')
        ldate = self.in_data.get('date', '')

        for entry in self.ustavki_entries:
            fname = os.path.basename(entry.get('file_path', ''))
            form  = entry.get('form_type', '')
            try:
                if form == 'new':
                    found_row, candidates = write_to_registry_new(entry, lnum, vx, ldate)
                    if found_row:
                        entry['registry_row'] = found_row
                        entry['status'] = 'реестр OK'
                        self._log(log, f"OK (стр {found_row})  {fname}")
                    else:
                        self._log(log, f"НЕ НАЙДЕНО в реестре: {fname}")
                        dlg = _RegistryNotFoundDialog(
                            self,
                            f"{entry.get('object_name','')} / {entry.get('dispatch_name','')}",
                            candidates,
                        )
                        if dlg.result_row:
                            write_registry_row_manual(dlg.result_row, lnum, vx, ldate)
                            entry['registry_row'] = dlg.result_row
                            entry['status'] = 'реестр OK (ручной)'
                            self._log(log, f"Записано вручную стр {dlg.result_row}")
                        else:
                            entry['status'] = 'пропущено'
                            self._log(log, "Пропущено")

                elif form == 'old':
                    found_row, sheet = write_to_summary_old(entry, lnum, vx)
                    if found_row:
                        entry['registry_row'] = found_row
                        entry['status'] = 'сводная OK'
                        self._log(log, f"OK лист {sheet} стр {found_row}  {fname}")
                    else:
                        entry['status'] = 'не найдено в сводной'
                        self._log(log, f"НЕ НАЙДЕНО в сводной: {fname}")
                else:
                    self._log(log, f"Форма неизвестна — пропуск: {fname}")

            except Exception as exc:
                entry['status'] = 'ошибка реестра'
                self._log(log, f"ОШИБКА  {fname}: {exc}")
        self._refresh_tree()


    # ─── Шаг 4: раскладка файлов ──────────────────────────────────────────

    def _find_archive_candidates_all(self):
        if not self._archive_tree:
            return
        self._archive_tree.delete(*self._archive_tree.get_children())
        for entry in self.ustavki_entries:
            obj    = entry.get('object_name', '')
            short  = match_object_to_short_name(obj)
            folder = find_object_exec_folder(short) if short else None
            current_dir, _ = find_current_and_archive_folders(folder) if folder else (None, None)
            candidates = find_archive_candidates(entry['file_path'], current_dir or '') if current_dir else []

            top_candidate = candidates[0] if candidates else ('','','')
            entry['archive_candidate'] = top_candidate[0]
            entry['_short_name']       = short
            entry['_current_dir']      = current_dir or ''
            entry['_archive_dir']      = find_current_and_archive_folders(folder)[1] if folder else ''

            self._archive_tree.insert('', 'end', values=(
                os.path.basename(entry['file_path']),
                short,
                top_candidate[1] if top_candidate else '',
                f"{top_candidate[2]:.2f}" if len(top_candidate) > 2 else '',
            ))

    def _on_archive_tree_dclick(self, event):
        """Двойной клик по строке показывает все кандидаты для выбора."""
        tree = self._archive_tree
        row_id = tree.identify_row(event.y)
        if not row_id:
            return
        all_ids = list(tree.get_children())
        try:
            idx = all_ids.index(row_id)
        except ValueError:
            return
        if idx >= len(self.ustavki_entries):
            return
        entry = self.ustavki_entries[idx]
        current_dir = entry.get('_current_dir', '')
        if not current_dir:
            return
        candidates = find_archive_candidates(entry['file_path'], current_dir, top_n=10)
        if not candidates:
            messagebox.showinfo("Кандидаты", "Похожих файлов не найдено.", parent=self)
            return

        # Диалог выбора кандидата
        dlg = tk.Toplevel(self)
        dlg.title("Выбор архивной таблицы")
        dlg.grab_set()
        ttk.Label(dlg, text=f"Кандидаты на архивацию для:\n{os.path.basename(entry['file_path'])}",
                  wraplength=460, font=('','10','bold')).pack(padx=12, pady=(10,4))
        cols2 = ('name', 'score')
        tv = ttk.Treeview(dlg, columns=cols2, show='headings', height=8)
        tv.heading('name', text='Имя файла'); tv.column('name', width=340)
        tv.heading('score', text='Схожесть');  tv.column('score', width=70)
        for cpath, cname, cscore in candidates:
            tv.insert('', 'end', iid=cpath, values=(cname, f"{cscore:.2f}"))
        tv.pack(padx=12, fill='both', expand=True)

        def _pick():
            sel = tv.selection()
            if sel:
                chosen_path = sel[0]
                entry['archive_candidate'] = chosen_path
                vals = list(tree.item(row_id, 'values'))
                vals[2] = os.path.basename(chosen_path)
                tree.item(row_id, values=vals)
            dlg.destroy()

        ttk.Button(dlg, text="Выбрать", command=_pick).pack(pady=8)
        dlg.transient(self)
        dlg.wait_window()

    def _move_files_all(self):
        moved, errors = 0, 0
        for entry in self.ustavki_entries:
            current_dir = entry.get('_current_dir', '')
            archive_dir = entry.get('_archive_dir', '')
            if not current_dir or not archive_dir:
                continue
            try:
                new_path = move_table_files(entry, archive_dir, current_dir)
                if new_path:
                    entry['current_path'] = new_path
                    entry['status'] = 'разложено'
                    moved += 1
            except Exception as exc:
                errors += 1
                messagebox.showerror("Ошибка раскладки",
                    f"{os.path.basename(entry['file_path'])}: {exc}", parent=self)
        self._refresh_tree()
        messagebox.showinfo("Раскладка завершена",
            f"Разложено: {moved}  Ошибок: {errors}", parent=self)

    # ─── Шаг 5: изменения (синие строки) ─────────────────────────────────

    def _create_changes_report(self):
        log = '5 Изменения'
        if not HAS_DOCX:
            self._log(log, "ОШИБКА: python-docx не установлен")
            return
        # Определяем путь для отчёта
        default_name = f"Изменения_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        out_path = filedialog.asksaveasfilename(
            title="Сохранить отчёт изменений",
            initialdir=os.path.expanduser('~'),
            initialfile=default_name,
            defaultextension='.docx',
            filetypes=[("Word файлы", "*.docx")],
        )
        if not out_path:
            return
        try:
            generate_changes_report(self.ustavki_entries, out_path)
            self._log(log, f"Отчёт создан: {out_path}")
        except Exception as exc:
            self._log(log, f"ОШИБКА: {exc}")

    # ─── Шаг 6: карты уставок ─────────────────────────────────────────────

    def _update_visio_maps_all(self):
        log = '6 Карты'
        for entry in self.ustavki_entries:
            short = match_object_to_short_name(entry.get('object_name', ''))
            if not short:
                self._log(log, f"Объект не распознан: {entry.get('object_name','')}")
                continue
            visio_path = os.path.join(MAPS_FOLDER, short + '.vsdx')
            if not os.path.exists(visio_path):
                visio_path = os.path.join(MAPS_FOLDER, short + '.vsd')
            if not os.path.exists(visio_path):
                self._log(log, f"Visio не найден для: {short}")
                continue
            old_path = entry.get('archive_candidate', '')
            new_path = entry.get('current_path', entry.get('file_path', ''))
            table_num = entry.get('table_number', '')
            ok, msg = update_visio_map(visio_path, old_path, new_path, table_num)
            self._log(log, f"{'OK' if ok else 'ERR'}  {short}: {msg}")

    # ─── Шаг 7: ДЭБ ───────────────────────────────────────────────────────

    def _upload_deb_all(self):
        log = '7 ДЭБ'
        if not HAS_SELENIUM or not HAS_PYAUTOGUI:
            self._log(log, "ОШИБКА: необходимы selenium и pyautogui")
            return
        for entry in self.ustavki_entries:
            short = match_object_to_short_name(entry.get('object_name', ''))
            visio_path = os.path.join(MAPS_FOLDER, short + '.vsdx') if short else ''
            pdf_path   = os.path.join(MAPS_PDF_FOLDER, short + '.pdf') if short else ''
            dispatch   = entry.get('dispatch_name', '')
            ok, msg    = upload_to_deb_entry(dispatch, visio_path, pdf_path)
            self._log(log, f"{'OK' if ok else 'ERR'}  {dispatch}: {msg}")

    # ─── Центровка окна ───────────────────────────────────────────────────

    def _center_window(self):
        self.update_idletasks()
        w, h = self.winfo_width(), self.winfo_height()
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        self.geometry(f"{max(w,900)}x{max(h,600)}+{(sw-max(w,900))//2}+{(sh-max(h,600))//2}")


# ── Точка входа ───────────────────────────────────────────────────────────────

def main():
    missing = []
    if not HAS_WIN32:
        missing.append("pywin32")
    if not HAS_OPENPYXL:
        missing.append("openpyxl")

    if missing:
        root = tk.Tk()
        root.withdraw()
        messagebox.showerror(
            "Не хватает зависимостей",
            "Не установлены библиотеки:\n  " + "\n  ".join(missing) +
            "\n\nПересоберите .exe с нужными зависимостями.",
        )
        root.destroy()
        sys.exit(1)

    if not HAS_DOCX:
        root = tk.Tk()
        root.withdraw()
        messagebox.showwarning(
            "python-docx не установлен",
            "Функции обработки таблиц уставок будут недоступны.\n"
            "Установите: pip install python-docx",
        )
        root.destroy()

    app = RegistrationApp()
    app.mainloop()


if __name__ == '__main__':
    main()
