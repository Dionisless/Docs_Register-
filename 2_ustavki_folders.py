#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
2_ustavki_folders.py — Обработка таблиц уставок: раскладка по папкам
=====================================================================
Выполняет шаги 0–4 цикла обработки таблиц уставок:
  0. Добавить файлы .docx
  1. Парсировать данные (объект, № таблицы, форма и т.д.)
  2. Записать «Уставки выставлены» в каждый .docx
  3. Записать в реестры (Реестр.xlsx / Регистрация.xlsx)
  4. Найти кандидатов на архив и разложить файлы по папкам

ВХОДНЫЕ ДАННЫЕ:
  - Файлы .docx (drag-and-drop или через диалог)
  - Данные письма из session_data.json → "letter" → "in_data"
    (сгенерированные программой 1_letter_register.py)

ИНТЕРФЕЙС ДЛЯ СШИВАНИЯ:
  - session_data.json → ключ "ustavki": список UstavkiEntry
  - UstavkiFoldersApp.ustavki_entries  — публичный список записей
  - get_entries() → list[dict]  — вернуть результат для программы 3
"""

import os
import re
import sys
import shutil
import difflib
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime


def _resource_path(rel: str) -> str:
    base = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, rel)

from shared_lib import (
    REGISTRY_PATH, SUMMARY_PATH, USTAVKI_EXEC_BASE, USTAVKI_ARCHIVE_BASE,
    USTAVKI_REAL_ARCHIVE_BASE, MAPS_FOLDER, MAPS_PDF_FOLDER,
    OBJECT_SHORT_NAMES, EMPTY_USTAVKI_ENTRY,
    sanitize_for_filename, parse_date, fmt_date_ymd_underscore,
    match_object_to_short_name, get_object_short_name_from_path,
    find_object_exec_folder, find_current_and_archive_folders,
    load_session, save_session,
)

# ── Зависимости ──────────────────────────────────────────────────────────────

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    HAS_DND = True
    _BASE_CLASS = TkinterDnD.Tk
except Exception:
    HAS_DND = False
    _BASE_CLASS = tk.Tk

try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.webdriver.chrome.options import Options as ChromeOptions
    from selenium.webdriver.chrome.service import Service as ChromeService
    HAS_SELENIUM = True
except ImportError:
    HAS_SELENIUM = False

try:
    import pyautogui
    HAS_PYAUTOGUI = True
except ImportError:
    HAS_PYAUTOGUI = False

try:
    import win32clipboard
    HAS_WIN32CLIP = True
except ImportError:
    HAS_WIN32CLIP = False

# ── Константы ─────────────────────────────────────────────────────────────────

DEB_BASE_URL = "https://pri-mdeb.oduvs.so"
DEB_MAPS_URL = ("https://pri-mdeb.oduvs.so/?sid=02ab815f-a54e-42a5-8a88-36dee8a5af2e"
                "&DataAreaId=1b6fecd6-f813-47ac-aa88-de4f67b7a1ac")

# ── Парсинг таблиц .docx ─────────────────────────────────────────────────────

def _parse_table_number_str(table_num: str) -> tuple | None:
    """
    Разбирает строку вида 'YY-NNN' или 'ПРДУ-РЗ-YY-NNN'.
    Возвращает (year_2digit, number) или None.
    """
    m = re.search(r'(\d{2})-(\d+)', table_num)
    if m:
        return int(m.group(1)), int(m.group(2))
    return None


def _is_old_form_by_number(table_num: str) -> bool:
    """
    Старая форма: год (YY) <= 25 И номер (NNN) <= 470.
    Примеры: 25-450 → старая, 24-1000 → старая, 25-471 → новая, 26-1 → новая.
    """
    parsed = _parse_table_number_str(table_num)
    if parsed is None:
        return False
    year, num = parsed
    return year <= 25 and num <= 470


def _extract_table_num_raw(doc_path: str, doc=None) -> str:
    """
    Пытается извлечь строку номера таблицы (YY-NNN) из:
    1) имени файла  (паттерн -YY-NNN( или _YY-NNN_)
    2) первого абзаца документа (ПРДУ-РЗ-YY-NNN)
    """
    fname = os.path.basename(doc_path)
    m = re.search(r'[-_](\d{2}-\d+)[(\s_]', fname)
    if m:
        return m.group(1)
    # Из абзацев — ПРДУ-РЗ-YY-NNN
    if doc is None:
        try:
            doc = DocxDocument(doc_path)
        except Exception:
            return ''
    for p in doc.paragraphs[:5]:
        m2 = re.search(r'ПРДУ-РЗ-(\d{2}-\d+)', p.text)
        if m2:
            return m2.group(1)
    return ''


def detect_table_form(doc_path: str) -> str:
    """
    Определяет форму таблицы уставок.
    Если удаётся извлечь номер — решение по правилу:
      старая: год(YY) <= 25 И номер(NNN) <= 470
      новая: иначе
    Если номер не найден — структурная эвристика (fallback).
    """
    try:
        doc = DocxDocument(doc_path)
    except Exception:
        return 'unknown'

    num_raw = _extract_table_num_raw(doc_path, doc)
    if num_raw:
        return 'old' if _is_old_form_by_number(num_raw) else 'new'

    # Fallback: структурная проверка
    if not doc.tables:
        return 'unknown'
    t = doc.tables[0]
    if not t.rows:
        return 'unknown'
    first_cell = t.rows[0].cells[0].text.strip()
    if len(t.columns) >= 3 and first_cell.isdigit():
        return 'old'
    return 'new'


def parse_ustavki_table(doc_path: str) -> dict:
    doc = DocxDocument(doc_path)
    fname = os.path.basename(doc_path)
    form = detect_table_form(doc_path)
    result = {
        'form_type': form, 'object_name': '', 'dispatch_name': '',
        'table_number': '', 'outgoing_letter': '', 'issue_reason': '',
    }
    if form == 'old':
        if doc.paragraphs:
            m = re.search(r'ПРДУ-РЗ-(\d{2}-\d+)', doc.paragraphs[0].text)
            if m:
                result['table_number'] = m.group(1)
        t = doc.tables[0]
        if len(t.rows) > 0:
            result['object_name'] = t.rows[0].cells[2].text.strip()
        if len(t.rows) > 1:
            result['dispatch_name'] = t.rows[1].cells[2].text.strip()
        for p in doc.paragraphs:
            if 'уставки выданы' in p.text.lower():
                after_colon = p.text.split(':', 1)[-1].strip()
                m2 = re.match(r'([\S]+)', after_colon)
                if m2:
                    result['outgoing_letter'] = m2.group(1)
                break
    else:
        m = re.search(r'-(\d{2}-\d+)\(', fname)
        if m:
            result['table_number'] = m.group(1)
        if doc.tables:
            t = doc.tables[0]
            if len(t.rows) > 0:
                result['object_name'] = t.rows[0].cells[1].text.strip()
            if len(t.rows) > 1:
                result['dispatch_name'] = t.rows[1].cells[1].text.strip()
            for row in t.rows:
                if row.cells and 'причина' in row.cells[0].text.lower():
                    result['issue_reason'] = row.cells[1].text.strip() if len(row.cells) > 1 else ''
                    break
        for p in doc.paragraphs:
            txt = p.text
            if 'к письму' in txt.lower() or re.search(r'№\s*[РрP]45', txt):
                m2 = re.search(r'№\s*([\w\-\.\/]+)\s*$', txt.strip())
                if m2:
                    result['outgoing_letter'] = m2.group(1)
                    break
    return result


# ── Word: запись «Уставки выставлены» ────────────────────────────────────────

def _levenshtein(s1: str, s2: str) -> int:
    if len(s1) < len(s2):
        return _levenshtein(s2, s1)
    if not s2:
        return len(s1)
    prev = list(range(len(s2) + 1))
    for i, c1 in enumerate(s1):
        curr = [i + 1]
        for j, c2 in enumerate(s2):
            curr.append(min(prev[j + 1] + 1, curr[j] + 1, prev[j] + (c1 != c2)))
        prev = curr
    return prev[-1]


def _find_issued_paragraph(doc):
    target_exact = 'уставки выставлены'
    target_fuzzy = 'таблицы выданы'
    for idx, p in enumerate(doc.paragraphs):
        if target_exact in p.text.lower():
            return p, idx
    win_len = len(target_fuzzy)
    for idx, p in enumerate(doc.paragraphs):
        low = p.text.lower()
        for i in range(len(low) - win_len + 1):
            if _levenshtein(low[i:i + win_len], target_fuzzy) <= 2:
                return p, idx
    return None, -1


def _add_hyperlink_run(paragraph, url: str, text: str):
    try:
        part = paragraph.part
        r_id = part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True,
        )
        hl = OxmlElement('w:hyperlink')
        hl.set(qn('r:id'), r_id)
        run_el = OxmlElement('w:r')
        rpr = OxmlElement('w:rPr')
        u_el = OxmlElement('w:u'); u_el.set(qn('w:val'), 'single')
        color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), '0563C1')
        rpr.append(u_el); rpr.append(color_el)
        run_el.append(rpr)
        t_el = OxmlElement('w:t'); t_el.text = text
        run_el.append(t_el)
        hl.append(run_el)
        paragraph._element.append(hl)
    except Exception:
        paragraph.add_run(text)


def write_issued_to_doc(doc_path: str, incoming_letter_num: str,
                        incoming_num: str, letter_date_str: str,
                        hyperlink_path: str = '') -> bool:
    doc = DocxDocument(doc_path)
    para, _ = _find_issued_paragraph(doc)
    if para is None:
        return False
    date_str = fmt_date_ymd_underscore(letter_date_str)
    label = f"{incoming_letter_num} {incoming_num} от {date_str}"
    if hyperlink_path:
        para.add_run(" ")
        _add_hyperlink_run(para, hyperlink_path, label)
    else:
        para.add_run(f" {label}")
    doc.save(doc_path)
    return True


# ── Реестр (новая форма) ──────────────────────────────────────────────────────

def _norm(v) -> str:
    if v is None: return ''
    return str(v).strip().lower()


def write_to_registry_new(entry: dict, incoming_letter_num: str,
                          incoming_num: str, letter_date) -> tuple:
    wb = openpyxl.load_workbook(REGISTRY_PATH)
    ws = wb['Реестр']
    obj   = _norm(entry.get('object_name', ''))
    disp  = _norm(entry.get('dispatch_name', ''))
    letter_out = _norm(entry.get('outgoing_letter', ''))
    found_row, candidates = None, []
    for row in range(ws.max_row, 2, -1):
        r_obj  = _norm(ws.cell(row, 2).value)
        r_disp = _norm(ws.cell(row, 3).value)
        r_out  = _norm(ws.cell(row, 4).value)
        mc = ((r_obj == obj and bool(obj)) +
              (r_disp == disp and bool(disp)) +
              (r_out == letter_out and bool(letter_out)))
        if mc == 3:
            found_row = row; break
        elif mc >= 2:
            candidates.append({
                'row': row, 'object': ws.cell(row, 2).value,
                'dispatch': ws.cell(row, 3).value,
                'outgoing': ws.cell(row, 4).value, 'matches': mc,
            })
    if found_row:
        _write_registry_row(ws, found_row, incoming_letter_num, incoming_num, letter_date)
        wb.save(REGISTRY_PATH)
    candidates.sort(key=lambda x: x['matches'], reverse=True)
    return found_row, candidates


def write_registry_row_manual(row: int, incoming_letter_num: str,
                               incoming_num: str, letter_date):
    wb = openpyxl.load_workbook(REGISTRY_PATH)
    ws = wb['Реестр']
    _write_registry_row(ws, row, incoming_letter_num, incoming_num, letter_date)
    wb.save(REGISTRY_PATH)


def _write_registry_row(ws, row: int, incoming_letter_num: str,
                        incoming_num: str, letter_date):
    ws.cell(row, 9).value  = incoming_letter_num
    ws.cell(row, 10).value = incoming_num
    if letter_date:
        if isinstance(letter_date, datetime):
            ws.cell(row, 11).value = letter_date
        else:
            dt = parse_date(str(letter_date))
            ws.cell(row, 11).value = dt if dt else str(letter_date)


# ── Регистрация (старая форма) ────────────────────────────────────────────────

def write_to_summary_old(entry: dict, incoming_letter_num: str,
                         incoming_num: str) -> tuple:
    table_num = entry.get('table_number', '')
    if not table_num:
        return None, ''
    m = re.match(r'(\d{2})-', table_num)
    year_sheet = ('20' + m.group(1)) if m else None
    wb = openpyxl.load_workbook(SUMMARY_PATH)
    service = {'служебный', 'шаблон', 'template'}
    year_sheets = [s for s in wb.sheetnames if s.lower() not in service]
    ws = None
    if year_sheet and year_sheet in wb.sheetnames:
        ws = wb[year_sheet]
    elif year_sheets:
        ws = wb[year_sheets[-1]]
    else:
        return None, ''
    num_col = exec_col = None
    for col in range(1, ws.max_column + 1):
        h = _norm(ws.cell(1, col).value)
        if '№' in h and ('таблиц' in h or 'задан' in h):
            num_col = col
        elif 'исполнен' in h and 'письмо' in h:
            exec_col = col
    if not num_col or not exec_col:
        return None, ws.title
    tnum_low = table_num.strip().lower()
    for row in range(ws.max_row, 1, -1):
        cell_val = _norm(ws.cell(row, num_col).value)
        if tnum_low in cell_val or cell_val == tnum_low:
            ws.cell(row, exec_col).value = f"{incoming_letter_num} вх-{incoming_num}"
            wb.save(SUMMARY_PATH)
            return row, ws.title
    return None, ws.title


# ── Файловые операции ─────────────────────────────────────────────────────────

def _strip_date_tail(s: str) -> str:
    """Убирает хвост вида (что-то_ДД_ММ_ГГГГ) в конце имени файла."""
    return re.sub(r'\s*\([^)]*\d{2}_\d{2}_\d{4}\)\s*$', '', s).strip()


def _normalize_name(s: str) -> str:
    return re.sub(r'\s+', ' ', s.lower().replace('ё', 'е').replace('-', ' ')).strip()


def _get_archive_object_folder(file_path: str) -> str:
    """
    Возвращает путь к папке объекта в Таблицы уставок РЗА.
    Берёт имя объекта из родительской папки обрабатываемого файла.
    Пример:
      file_path = ...\\Таблицы для исполнения РЗА\\Волна\\Текущие\\файл.docx
      → возвращает  ...\\Таблицы уставок РЗА\\Волна
    """
    # Имя папки-объекта = папка-дедушка или папка-родитель файла
    # Ищем ближайшую к файлу папку, имя которой есть в OBJECT_SHORT_NAMES
    obj_short = get_object_short_name_from_path(file_path)
    if not obj_short:
        # Fallback: просто берём родительскую папку
        obj_short = os.path.basename(os.path.dirname(file_path))
    if not obj_short:
        return ''
    archive_folder = os.path.join(USTAVKI_ARCHIVE_BASE, obj_short)
    return archive_folder


def find_archive_candidates_by_filename(new_filepath: str, top_n: int = 5) -> list:
    """
    Вариант А: ищет кандидата на архив по схожести ИМЕНИ ФАЙЛА.
    Смотрит в папку \\Таблицы уставок РЗА\\Объект\\.
    Возвращает список (path, name, score).
    """
    archive_folder = _get_archive_object_folder(new_filepath)
    if not archive_folder or not os.path.isdir(archive_folder):
        return []
    new_stem = _normalize_name(_strip_date_tail(os.path.splitext(os.path.basename(new_filepath))[0]))
    candidates = []
    try:
        for entry in os.scandir(archive_folder):
            if not entry.is_file():
                continue
            stem = _normalize_name(_strip_date_tail(os.path.splitext(entry.name)[0]))
            score = difflib.SequenceMatcher(None, new_stem, stem).ratio()
            candidates.append((entry.path, entry.name, score))
    except OSError:
        pass
    candidates.sort(key=lambda x: x[2], reverse=True)
    return candidates[:top_n]


def find_archive_candidates_by_dispatch(new_filepath: str, dispatch_name: str,
                                        top_n: int = 5) -> list:
    """
    Вариант Б: ищет кандидата на архив по схожести ДИСПЕТЧЕРСКОГО НАИМЕНОВАНИЯ.
    Парсирует dispatch_name из каждого файла в \\Таблицы уставок РЗА\\Объект\\
    и выбирает наиболее похожее на dispatch_name обрабатываемой таблицы.
    Возвращает список (path, name, score, parsed_dispatch).
    """
    if not HAS_DOCX:
        return []
    archive_folder = _get_archive_object_folder(new_filepath)
    if not archive_folder or not os.path.isdir(archive_folder):
        return []
    norm_dispatch = _normalize_name(dispatch_name)
    candidates = []
    try:
        for entry in os.scandir(archive_folder):
            if not entry.is_file() or not entry.name.lower().endswith(('.docx', '.doc')):
                continue
            try:
                parsed = parse_ustavki_table(entry.path)
                candidate_dispatch = parsed.get('dispatch_name', '') or parsed.get('object_name', '')
                norm_cd = _normalize_name(candidate_dispatch)
                score = difflib.SequenceMatcher(None, norm_dispatch, norm_cd).ratio()
                candidates.append((entry.path, entry.name, score, candidate_dispatch))
            except Exception:
                # Если файл нельзя прочитать — пропускаем
                continue
    except OSError:
        pass
    candidates.sort(key=lambda x: x[2], reverse=True)
    return candidates[:top_n]


def find_archive_candidates(new_filepath: str, current_dir: str, top_n: int = 5) -> list:
    """Обратная совместимость: поиск по имени файла в указанной папке."""
    if not current_dir or not os.path.isdir(current_dir):
        return []
    new_stem = _normalize_name(_strip_date_tail(os.path.splitext(os.path.basename(new_filepath))[0]))
    candidates = []
    try:
        for entry in os.scandir(current_dir):
            if not entry.is_file():
                continue
            stem = _normalize_name(_strip_date_tail(os.path.splitext(entry.name)[0]))
            score = difflib.SequenceMatcher(None, new_stem, stem).ratio()
            candidates.append((entry.path, entry.name, score))
    except OSError:
        pass
    candidates.sort(key=lambda x: x[2], reverse=True)
    return candidates[:top_n]


def move_table_files(entry: dict, archive_dir: str, current_dir: str) -> str:
    os.makedirs(archive_dir, exist_ok=True)
    os.makedirs(current_dir, exist_ok=True)
    archive_src = entry.get('archive_candidate', '')
    if archive_src and os.path.exists(archive_src):
        dest = os.path.join(archive_dir, os.path.basename(archive_src))
        shutil.move(archive_src, dest)
    src = entry.get('file_path', '')
    if src and os.path.exists(src):
        dest_new = os.path.join(current_dir, os.path.basename(src))
        shutil.move(src, dest_new)
        return dest_new
    return ''


# ── Шаг 5: синие строки и отчёт изменений ────────────────────────────────────

_BLUE_COLORS = {c.lower() for c in [
    '4472C4','5B9BD5','2E74B5','0070C0','00B0F0','1F3864',
    '2F5597','1155CC','0000FF','538DD5','4F81BD','44546A',
]}


def _run_has_blue_text(run_element) -> bool:
    try:
        from lxml import etree
    except ImportError:
        return False
    run_text = ''.join(t.text or '' for t in run_element.iter()
                       if t.tag.endswith('}t') or t.tag == 't')
    if not re.search(r'[a-zA-Zа-яА-ЯёЁ0-9]', run_text):
        return False
    xml_str = etree.tostring(run_element).decode('utf-8', errors='ignore').lower()
    for bc in _BLUE_COLORS:
        if bc in xml_str:
            return True
    return False


_BLUE_BOILERPLATE_FRAGMENTS = [
    'параметры настройки, измененные относительно предыдущего задания',
    'выделены цветом и жирным шрифтом',
]


def _is_boilerplate_row(cells_text: list) -> bool:
    joined = ' '.join(cells_text).lower()
    return any(f in joined for f in _BLUE_BOILERPLATE_FRAGMENTS)


def extract_blue_rows_from_doc(doc_path: str) -> list:
    if not HAS_DOCX:
        return []
    try:
        from lxml import etree
    except ImportError:
        return []
    doc = DocxDocument(doc_path)
    W_R = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'
    result = []
    for table in doc.tables:
        for row in table.rows:
            has_blue = False
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run_el in para._element.iter(W_R):
                        if _run_has_blue_text(run_el):
                            has_blue = True
                            break
                    if has_blue:
                        break
                if has_blue:
                    break
            if has_blue:
                cells_text = [c.text.strip() for c in row.cells]
                if not _is_boilerplate_row(cells_text):
                    result.append(cells_text)
    return result


def generate_changes_report(entries: list, output_path: str):
    if not HAS_DOCX:
        raise RuntimeError("python-docx не установлен")
    doc = DocxDocument()
    doc.add_heading('Сводка изменений таблиц уставок', level=1)
    for i, entry in enumerate(entries, 1):
        doc.add_heading(
            f"Таблица {i}: {entry.get('dispatch_name','?')} — {entry.get('table_number','?')}",
            level=2)
        info_lines = [
            f"Объект: {entry.get('object_name','')}",
            f"Исх. письмо: {entry.get('outgoing_letter','')}",
        ]
        if entry.get('form_type') == 'new' and entry.get('issue_reason'):
            info_lines.append(f"Причина выдачи: {entry.get('issue_reason')}")
        doc.add_paragraph('\n'.join(info_lines))
        blue_rows = extract_blue_rows_from_doc(entry.get('file_path', ''))
        if not blue_rows:
            doc.add_paragraph('Изменений (синий цвет) не обнаружено.')
            continue
        num_cols = max(len(r) for r in blue_rows)
        tbl = doc.add_table(rows=len(blue_rows), cols=num_cols)
        tbl.style = 'Table Grid'
        for ri, row_data in enumerate(blue_rows):
            for ci, cell_text in enumerate(row_data):
                if ci < num_cols:
                    tbl.cell(ri, ci).text = cell_text
    doc.save(output_path)
    try:
        os.startfile(output_path)
    except Exception:
        pass


# ── Шаг 6: Visio ──────────────────────────────────────────────────────────────

def _visio_iter_shapes(shapes_coll):
    """Рекурсивно обходит все фигуры включая sub-shapes внутри групп."""
    count = 0
    try:
        count = shapes_coll.Count
    except Exception:
        return
    for i in range(1, count + 1):
        try:
            shp = shapes_coll.Item(i)
        except Exception:
            continue
        yield shp
        try:
            if shp.Shapes.Count > 0:
                yield from _visio_iter_shapes(shp.Shapes)
        except Exception:
            pass


def _win_basename(path: str) -> str:
    """Возвращает имя файла из Windows-пути (работает на любой ОС)."""
    return path.replace('/', '\\').split('\\')[-1]


def _extract_table_num_from_filename(path: str) -> str:
    """
    Извлекает номер таблицы (YY-NNN) из имени файла.
    Пример: 'ВТЭЦ-2 ШСМВ-110-МТЗ ТЗНП-23-220.docx' → '23-220'
    Берёт ПОСЛЕДНЕЕ вхождение шаблона \\d{2}-\\d+ в стебле имени.
    """
    stem = os.path.splitext(_win_basename(path))[0]
    matches = re.findall(r'\d{2}-\d+', stem)
    return matches[-1] if matches else ''


def update_visio_map(visio_path: str, old_table_path: str,
                     new_table_path: str, new_table_number: str,
                     pdf_folder: str = '') -> tuple:
    """
    Обновляет гиперссылку в карте уставок (Visio):
      - Ищет гиперссылку по имени файла (basename) старой таблицы
        (hl.Address хранится как относительный путь → точное сравнение
         абсолютных путей всегда даёт 0; ищем по имени файла)
      - Заменяет Address на абсолютный путь к новой таблице
      - В Description записывает старое имя файла (история замены)
      - В shape.Text заменяет YY-NNN старой таблицы на номер новой
      - Сохраняет .vsdx/.vsd и экспортирует PDF
    """
    try:
        import win32com.client as win32
    except ImportError:
        return False, "win32com не доступен (нужен pywin32)"
    if not os.path.exists(visio_path):
        return False, f"Файл не найден: {visio_path}"
    if not new_table_path:
        return False, "Не указан новый путь к файлу таблицы"

    # Базовые имена для сравнения (без учёта регистра)
    old_basename = _win_basename(old_table_path).lower() if old_table_path else ''
    old_abs_lower = old_table_path.lower().replace('/', '\\') if old_table_path else ''
    # Номер таблицы в старом файле (для замены в shape.Text)
    old_yynum = _extract_table_num_from_filename(old_table_path) if old_table_path else ''

    visio = None
    try:
        visio = win32.Dispatch('Visio.Application')
        visio.Visible = False
        doc = visio.Documents.Open(os.path.abspath(visio_path))
        replaced = 0
        for page in doc.Pages:
            for shape in _visio_iter_shapes(page.Shapes):
                try:
                    hl_count = shape.Hyperlinks.Count
                except Exception:
                    continue
                for i in range(1, hl_count + 1):
                    try:
                        hl = shape.Hyperlinks.Item(i)
                        addr = hl.Address or ''
                        if not old_basename:
                            continue
                        addr_norm = addr.replace('/', '\\')
                        addr_basename = _win_basename(addr_norm).lower()

                        # Сравниваем по имени файла (основной способ —
                        # Visio хранит относительные пути)
                        matched = (old_basename == addr_basename)
                        # Запасной вариант: если Address уже абсолютный
                        if not matched and old_abs_lower:
                            matched = old_abs_lower in addr_norm.lower()

                        if not matched:
                            continue

                        # Обновляем гиперссылку
                        hl.Address = new_table_path
                        # Description: старое имя файла как история замены
                        hl.Description = _win_basename(old_table_path) if old_table_path else ''

                        # Обновляем видимый текст фигуры:
                        # заменяем YY-NNN старой таблицы на номер новой
                        if new_table_number and old_yynum:
                            try:
                                old_text = shape.Text
                                if old_yynum in old_text:
                                    shape.Text = old_text.replace(old_yynum, new_table_number)
                            except Exception:
                                pass

                        replaced += 1
                    except Exception:
                        continue
        doc.Save()
        stem = os.path.splitext(os.path.basename(visio_path))[0]
        out_dir = pdf_folder if pdf_folder else MAPS_PDF_FOLDER
        pdf_path = os.path.join(out_dir, stem + '.pdf')
        os.makedirs(out_dir, exist_ok=True)
        pdf_exported = False
        pdf_err = ''
        for args in [(1, pdf_path, 0, 0), (1, pdf_path, 0), (1, pdf_path)]:
            try:
                doc.ExportAsFixedFormat(*args)
                pdf_exported = True
                break
            except Exception as e:
                pdf_err = str(e)
        doc.Close()
        if pdf_exported:
            return True, f"Заменено ссылок: {replaced}, PDF: {pdf_path}"
        else:
            return True, f"Заменено ссылок: {replaced} (PDF не экспортирован: {pdf_err})"
    except Exception as exc:
        return False, str(exc)
    finally:
        try:
            if visio:
                visio.Quit()
        except Exception:
            pass


# ── ДЭБ: загрузка карт уставок ───────────────────────────────────────────────

def _find_yandex_binary() -> str:
    """Ищет исполняемый файл Яндекс Браузера на Windows."""
    candidates = [
        os.path.join(os.environ.get('LOCALAPPDATA', ''),
                     r'Yandex\YandexBrowser\Application\browser.exe'),
        os.path.join(os.environ.get('PROGRAMFILES', ''),
                     r'Yandex\YandexBrowser\Application\browser.exe'),
        os.path.join(os.environ.get('PROGRAMFILES(X86)', ''),
                     r'Yandex\YandexBrowser\Application\browser.exe'),
    ]
    for c in candidates:
        if c and os.path.exists(c):
            return c
    return ''


def _find_webdriver(driver_path: str = '') -> 'ChromeService | None':
    """
    Ищет yandexdriver.exe / chromedriver.exe:
    1. В явно заданном пути driver_path
    2. Рядом с exe (или рядом со скриптом при запуске без сборки)
    3. В PATH (возвращает None — selenium найдёт сам)
    """
    if driver_path and os.path.exists(driver_path):
        return ChromeService(driver_path)
    exe_dir = (os.path.dirname(sys.executable)
               if getattr(sys, 'frozen', False)
               else os.path.dirname(os.path.abspath(__file__)))
    for name in ('yandexdriver.exe', 'chromedriver.exe',
                 'yandexdriver', 'chromedriver'):
        cand = os.path.join(exe_dir, name)
        if os.path.exists(cand):
            return ChromeService(cand)
    return None


def _clipboard_set(text: str):
    """Кладёт текст в буфер обмена с поддержкой Unicode (через win32clipboard или tkinter)."""
    if HAS_WIN32CLIP:
        win32clipboard.OpenClipboard()
        try:
            win32clipboard.EmptyClipboard()
            win32clipboard.SetClipboardText(text, win32clipboard.CF_UNICODETEXT)
        finally:
            win32clipboard.CloseClipboard()
    else:
        # Fallback через Tk — временное окно
        try:
            r = tk.Tk(); r.withdraw()
            r.clipboard_clear(); r.clipboard_append(text); r.update()
            r.after(300, r.destroy); r.mainloop()
        except Exception:
            pass


def upload_to_deb_entry(dispatch_name: str, visio_path: str, pdf_path: str,
                        driver_path: str = '') -> tuple:
    """
    Загружает файлы карты уставок в ДЭБ через Яндекс Браузер + Selenium.
    Возвращает (success: bool, message: str).

    driver_path — явный путь к yandexdriver.exe; если пустой — ищется
                  рядом с exe или в PATH.
    """
    if not HAS_SELENIUM:
        return False, "selenium не установлен (pip install selenium)"
    if not HAS_PYAUTOGUI:
        return False, "pyautogui не установлен (pip install pyautogui)"

    opts = ChromeOptions()
    yandex_bin = _find_yandex_binary()
    if yandex_bin:
        opts.binary_location = yandex_bin
    opts.add_argument('--no-sandbox')
    opts.add_argument('--disable-dev-shm-usage')
    # Не закрывать браузер при ошибке — удобно для отладки
    opts.add_experimental_option('detach', False)

    service = _find_webdriver(driver_path) if HAS_SELENIUM else None
    driver = None
    try:
        driver = (webdriver.Chrome(service=service, options=opts)
                  if service else webdriver.Chrome(options=opts))
        wait = WebDriverWait(driver, 30)

        # 1. Открываем каталог ДЭБ
        driver.get(DEB_MAPS_URL)
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, 'table tbody tr')))
        time.sleep(1.0)

        # 2. Ищем строку с нашим объектом
        rows = driver.find_elements(By.CSS_SELECTOR, 'table tbody tr')
        target_row = None
        dn_low = dispatch_name.lower()
        for row in rows:
            try:
                tds = row.find_elements(By.TAG_NAME, 'td')
                for td in tds:
                    if dn_low in td.text.lower():
                        target_row = row
                        break
            except Exception:
                pass
            if target_row:
                break

        if not target_row:
            driver.quit()
            return False, f"Объект не найден в ДЭБ: {dispatch_name}"

        # 3. Переходим по ссылке «Перейти по ссылке»
        link_btn = target_row.find_element(By.CSS_SELECTOR, 'a[title="Перейти по ссылке"]')
        href = link_btn.get_attribute('href') or ''
        if href.startswith('/'):
            card_url = DEB_BASE_URL + href
        elif href.startswith('http'):
            card_url = href
        else:
            card_url = DEB_BASE_URL + '/' + href.lstrip('./')
        driver.execute_script("window.location.href = arguments[0];", card_url)
        wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, '.edit-mode-button')))
        time.sleep(1.0)

        # 4. Нажимаем «Редактировать»
        driver.find_element(By.CSS_SELECTOR, 'button.edit-mode-button').click()
        time.sleep(1.5)

        # 5. Обрабатываем файловые строки (visio и pdf)
        file_items = driver.find_elements(By.CSS_SELECTOR, '.list-group-item')
        uploaded = []
        for item in file_items:
            try:
                filename_el = item.find_element(By.CSS_SELECTOR, 'a.dz-filename')
                fname_text = filename_el.text.lower()
                is_visio = any(ext in fname_text for ext in ('.vsd', '.vsdx'))
                is_pdf   = '.pdf' in fname_text
                upload_path = visio_path if is_visio else (pdf_path if is_pdf else None)
                if not upload_path or not os.path.exists(upload_path):
                    continue

                change_btn = item.find_element(By.CSS_SELECTOR, 'a.change-file-button')
                change_btn.click()
                time.sleep(0.8)

                # Подтверждение (бывает только для visio-файла)
                try:
                    confirm_btn = WebDriverWait(driver, 4).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, 'button.btn.btn-success'))
                    )
                    confirm_btn.click()
                    time.sleep(0.8)
                except Exception:
                    pass

                # Ввод пути в диалог проводника
                time.sleep(1.5)
                _clipboard_set(upload_path)
                time.sleep(0.3)
                pyautogui.hotkey('ctrl', 'a')
                time.sleep(0.1)
                pyautogui.hotkey('ctrl', 'v')
                time.sleep(0.3)
                pyautogui.press('enter')
                time.sleep(2.5)
                uploaded.append(os.path.basename(upload_path))

            except Exception:
                continue

        # 6. Сохранить
        save_btn = driver.find_element(
            By.CSS_SELECTOR, 'button.confirm-button.btn-outline-success')
        save_btn.click()
        time.sleep(2.0)
        driver.quit()
        return True, f"Загружено для «{dispatch_name}»: {', '.join(uploaded)}"

    except Exception as exc:
        try:
            if driver:
                driver.quit()
        except Exception:
            pass
        return False, str(exc)


# ── UI: диалог «Строка не найдена» ───────────────────────────────────────────

class _RegistryNotFoundDialog(tk.Toplevel):
    def __init__(self, parent, entry_desc: str, candidates: list):
        super().__init__(parent)
        self.title("Строка не найдена в реестре")
        self.resizable(True, False)
        self.grab_set()
        self.result_row = None

        ttk.Label(self, text=f"Не найдено: {entry_desc}", wraplength=480,
                  font=('', 10, 'bold')).pack(padx=12, pady=(10, 4))
        if candidates:
            ttk.Label(self, text="Похожие строки:").pack(padx=12, anchor='w')
            frame = ttk.Frame(self)
            frame.pack(fill='both', expand=True, padx=12, pady=4)
            cols = ('row', 'object', 'dispatch', 'outgoing')
            tv = ttk.Treeview(frame, columns=cols, show='headings', height=6)
            tv.heading('row',      text='Строка'); tv.column('row',      width=60)
            tv.heading('object',   text='Объект'); tv.column('object',   width=160)
            tv.heading('dispatch', text='Дисп.');  tv.column('dispatch', width=200)
            tv.heading('outgoing', text='Исх.');   tv.column('outgoing', width=120)
            for c in candidates:
                tv.insert('', 'end', values=(
                    c['row'], c.get('object',''), c.get('dispatch',''), c.get('outgoing','')))
            sb = ttk.Scrollbar(frame, orient='vertical', command=tv.yview)
            tv.configure(yscrollcommand=sb.set)
            tv.pack(side='left', fill='both', expand=True)
            sb.pack(side='right', fill='y')
            tv.bind('<<TreeviewSelect>>', lambda e: self._tv_select(tv))
            self._tv = tv

        row_frame = ttk.Frame(self)
        row_frame.pack(padx=12, pady=4)
        ttk.Label(row_frame, text="Строка вручную:").pack(side='left')
        self._row_var = tk.StringVar()
        ttk.Entry(row_frame, textvariable=self._row_var, width=8).pack(side='left', padx=4)

        btn_frame = ttk.Frame(self)
        btn_frame.pack(pady=8)
        ttk.Button(btn_frame, text="Принять",   command=self._on_accept).pack(side='left', padx=6)
        ttk.Button(btn_frame, text="Пропустить",command=self.destroy).pack(side='left', padx=6)
        self.transient(parent)
        self.wait_window()

    def _tv_select(self, tv):
        sel = tv.selection()
        if sel:
            val = tv.item(sel[0], 'values')
            self._row_var.set(val[0] if val else '')

    def _on_accept(self):
        try:
            self.result_row = int(self._row_var.get())
        except ValueError:
            messagebox.showwarning("Ошибка", "Введите целое число.", parent=self)
            return
        self.destroy()


# ── Главное окно ─────────────────────────────────────────────────────────────

class UstavkiFoldersApp(_BASE_CLASS):
    """
    Программа 2: Раскладка таблиц уставок по папкам.

    Публичный интерфейс:
      .ustavki_entries   list[dict]  — обрабатываемые записи
      .in_data           dict        — данные входящего письма
      .get_entries() → list[dict]    — результат для программы 3
    """

    def __init__(self):
        super().__init__()
        self.ustavki_entries: list = []
        self.in_data: dict = {}
        self._candidates_found: bool = False   # True после нажатия «Найти кандидатов»

        # Загрузить сессию
        session = load_session()
        if 'letter' in session:
            self.in_data = session['letter'].get('in_data', {})
        if 'ustavki' in session:
            self.ustavki_entries = session['ustavki']

        self.title("Таблицы уставок — Раскладка по папкам  v2")
        self.resizable(True, True)
        try:
            self.iconbitmap(_resource_path('icons/2_folders.ico'))
        except Exception:
            pass
        self._build_ui()
        self._center_window()

    # ── Публичный интерфейс ────────────────────────────────────────────────

    def get_entries(self) -> list:
        """Возвращает список записей для передачи в программу 3."""
        return [dict(e) for e in self.ustavki_entries]

    # ── UI ────────────────────────────────────────────────────────────────

    def _build_ui(self):
        root = ttk.Frame(self, padding=6)
        root.grid(row=0, column=0, sticky='nsew')
        self.columnconfigure(0, weight=1)
        self.rowconfigure(0, weight=1)
        root.columnconfigure(0, weight=1)
        root.rowconfigure(1, weight=1)

        # Полоска письма
        lf = ttk.LabelFrame(root, text="Письмо о выставлении (из сессии / вводить вручную)", padding=6)
        lf.grid(row=0, column=0, sticky='ew', pady=(0, 4))
        lf.columnconfigure(3, weight=1)
        ttk.Label(lf, text="№ вх:").grid(row=0, column=0, sticky='e', padx=(0,3))
        self._lu_vx = tk.StringVar(value=self.in_data.get('incoming_num',''))
        ttk.Entry(lf, textvariable=self._lu_vx, width=14).grid(row=0, column=1, sticky='w')
        ttk.Label(lf, text="№ письма:").grid(row=0, column=2, sticky='e', padx=(10,3))
        self._lu_letter = tk.StringVar(value=self.in_data.get('letter_num',''))
        ttk.Entry(lf, textvariable=self._lu_letter, width=26).grid(row=0, column=3, sticky='w')
        ttk.Label(lf, text="Дата:").grid(row=0, column=4, sticky='e', padx=(10,3))
        self._lu_date = tk.StringVar(value=self.in_data.get('date',''))
        ttk.Entry(lf, textvariable=self._lu_date, width=12).grid(row=0, column=5, sticky='w')
        ttk.Label(lf, text="Ссылка:").grid(row=0, column=6, sticky='e', padx=(10,3))
        self._lu_link = tk.StringVar(value=self.in_data.get('file_link',''))
        ttk.Entry(lf, textvariable=self._lu_link, width=30).grid(row=0, column=7, sticky='ew')

        # Шаги
        self._step_nb = ttk.Notebook(root)
        self._step_nb.grid(row=1, column=0, sticky='nsew', pady=4)

        step0 = ttk.Frame(self._step_nb, padding=8)
        step1 = ttk.Frame(self._step_nb, padding=8)
        step2 = ttk.Frame(self._step_nb, padding=8)
        step3 = ttk.Frame(self._step_nb, padding=8)
        step4 = ttk.Frame(self._step_nb, padding=8)
        step5 = ttk.Frame(self._step_nb, padding=8)
        step6 = ttk.Frame(self._step_nb, padding=8)
        step7 = ttk.Frame(self._step_nb, padding=8)
        self._step_nb.add(step0, text=" 0 Файлы ")
        self._step_nb.add(step1, text=" 1 Данные ")
        self._step_nb.add(step2, text=" 2 → .docx ")
        self._step_nb.add(step3, text=" 3 Реестры ")
        self._step_nb.add(step4, text=" 4 Раскладка ")
        self._step_nb.add(step5, text=" 5 Изменения ")
        self._step_nb.add(step6, text=" 6 Карты Visio ")
        self._step_nb.add(step7, text=" 7 ДЭБ ")

        self._build_step0(step0)
        self._build_step1(step1)
        self._build_step_log(step2,
            "2_write_issued",
            "Запишет «Уставки выставлены: ПИСЬМО вх-N от ДАТА» в каждый файл .docx\n"
            "Перед записью показывает что и куда будет вставлено — требует подтверждения.",
            "Записать в таблицы →", self._write_issued_all)
        self._build_step_log(step3,
            "3_registries",
            "Новая форма → Реестр таблиц уставок.xlsx\n"
            "Старая форма → Регистрация таблиц уставок.xlsx\n"
            "При неудаче — предложит выбрать строку вручную.",
            "Записать в реестры →", self._write_registries_all)
        self._build_step4(step4)
        self._build_step5(step5)
        self._build_step6(step6)
        self._build_step7(step7)

        # Нижняя панель
        bot = ttk.Frame(root)
        bot.grid(row=2, column=0, pady=4, sticky='ew')
        ttk.Button(bot, text="Загрузить сессию", command=self._load_session).pack(side='left', padx=4)
        ttk.Button(bot, text="Сохранить сессию", command=self._save_current_session).pack(side='left', padx=4)
        ttk.Button(bot, text="Закрыть", command=self.destroy).pack(side='right', padx=4)

    def _build_step0(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(2, weight=1)

        dnd_frame = ttk.LabelFrame(parent, text="Файлы .docx/.doc — перетащить или добавить кнопкой", padding=8)
        dnd_frame.grid(row=0, column=0, sticky='ew', pady=(0,6))
        ttk.Label(dnd_frame, text="Перетащите файлы сюда или нажмите «Добавить файлы»",
                  foreground='gray', justify='center').pack(pady=12)
        if HAS_DND:
            dnd_frame.drop_target_register(DND_FILES)
            dnd_frame.dnd_bind('<<Drop>>', self._on_dnd_drop)

        btn = ttk.Frame(parent)
        btn.grid(row=1, column=0, sticky='ew', pady=(0,4))
        ttk.Button(btn, text="Добавить файлы…", command=self._add_files).pack(side='left', padx=4)
        ttk.Button(btn, text="Удалить выбранные", command=self._remove_selected).pack(side='left', padx=4)
        ttk.Button(btn, text="→ Парсировать", command=self._go_parse).pack(side='right', padx=4)

        lf = ttk.LabelFrame(parent, text="Добавленные файлы", padding=4)
        lf.grid(row=2, column=0, sticky='nsew')
        lf.columnconfigure(0, weight=1)
        lf.rowconfigure(0, weight=1)
        sb = ttk.Scrollbar(lf, orient='vertical')
        self._lb = tk.Listbox(lf, yscrollcommand=sb.set, selectmode='extended', height=10)
        sb.configure(command=self._lb.yview)
        self._lb.grid(row=0, column=0, sticky='nsew')
        sb.grid(row=0, column=1, sticky='ns')
        self._refresh_lb()

    def _build_step1(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)
        tf = ttk.Frame(parent)
        tf.grid(row=0, column=0, sticky='nsew')
        tf.columnconfigure(0, weight=1)
        tf.rowconfigure(0, weight=1)
        cols = ('file','form','object','dispatch','table_num',
                'out_letter','letter_num','letter_date','status')
        self._tree = ttk.Treeview(tf, columns=cols, show='headings', height=14)
        widths = [160,55,160,200,90,140,140,100,90]
        heads  = ['Файл','Форма','Объект','Дисп. наим.',
                  '№ таблицы','Исх. письмо','Вх. письмо','Дата вх.','Статус']
        for col, head, w in zip(cols, heads, widths):
            self._tree.heading(col, text=head)
            self._tree.column(col, width=w, minwidth=40)
        vsb = ttk.Scrollbar(tf, orient='vertical',   command=self._tree.yview)
        hsb = ttk.Scrollbar(tf, orient='horizontal', command=self._tree.xview)
        self._tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self._tree.grid(row=0, column=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        hsb.grid(row=1, column=0, sticky='ew')
        self._tree.bind('<Double-1>', self._on_tree_dclick)
        btn = ttk.Frame(parent)
        btn.grid(row=1, column=0, sticky='ew', pady=(6,0))
        ttk.Button(btn, text="Парсировать все",   command=self._parse_all).pack(side='left', padx=4)
        ttk.Button(btn, text="Применить письмо",  command=self._apply_letter).pack(side='left', padx=4)
        ttk.Button(btn, text="Очистить список",   command=self._clear_all).pack(side='right', padx=4)

    def _build_step_log(self, parent, key, description, btn_text, btn_cmd):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)
        ttk.Label(parent, text=description, wraplength=700,
                  foreground='gray').grid(row=0, column=0, sticky='w', pady=(0,6))
        txt = tk.Text(parent, height=14, wrap='word', state='disabled', font=('Consolas', 9))
        sb = ttk.Scrollbar(parent, orient='vertical', command=txt.yview)
        txt.configure(yscrollcommand=sb.set)
        txt.grid(row=1, column=0, sticky='nsew')
        sb.grid(row=1, column=1, sticky='ns')
        ttk.Button(parent, text=btn_text, command=btn_cmd).grid(
            row=2, column=0, pady=(6,0), sticky='w')
        if not hasattr(self, '_logs'):
            self._logs = {}
        self._logs[key] = txt

    def _build_step4(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)
        ttk.Label(parent,
            text="Ищет кандидата на архив в папке \\\\Prim-fs-serv\\...\\Таблицы уставок РЗА\\Объект\\\n"
                 "по схожести имени файла. Двойной клик — выбрать другого кандидата из списка.",
            wraplength=700, foreground='gray').grid(row=0, column=0, sticky='w', pady=(0, 4))
        af = ttk.Frame(parent)
        af.grid(row=1, column=0, sticky='nsew')
        af.columnconfigure(0, weight=1)
        af.rowconfigure(0, weight=1)
        arc_cols = ('file', 'short', 'candidate', 'score', 'method')
        self._arc_tree = ttk.Treeview(af, columns=arc_cols, show='headings', height=12)
        self._arc_tree.heading('file',      text='Новый файл');        self._arc_tree.column('file',      width=180)
        self._arc_tree.heading('short',     text='Объект (папка)');    self._arc_tree.column('short',     width=110)
        self._arc_tree.heading('candidate', text='Кандидат на архив'); self._arc_tree.column('candidate', width=220)
        self._arc_tree.heading('score',     text='Схожесть');          self._arc_tree.column('score',     width=65)
        self._arc_tree.heading('method',    text='Метод');             self._arc_tree.column('method',    width=50)
        vsb = ttk.Scrollbar(af, orient='vertical',   command=self._arc_tree.yview)
        hsb = ttk.Scrollbar(af, orient='horizontal', command=self._arc_tree.xview)
        self._arc_tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        self._arc_tree.grid(row=0, column=0, sticky='nsew')
        vsb.grid(row=0, column=1, sticky='ns')
        hsb.grid(row=1, column=0, sticky='ew')
        self._arc_tree.bind('<Double-1>', self._on_arc_dclick)
        bf = ttk.Frame(parent)
        bf.grid(row=2, column=0, pady=(6, 0), sticky='w')
        ttk.Button(bf, text="Найти кандидатов на архив",
                   command=self._find_candidates_a).pack(side='left', padx=4)
        ttk.Button(bf, text="Разложить файлы",
                   command=self._move_files).pack(side='left', padx=4)

    # ── Логирование ───────────────────────────────────────────────────────

    def _log(self, key: str, text: str):
        widget = self._logs.get(key)
        if widget:
            widget.configure(state='normal')
            widget.insert('end', text + '\n')
            widget.see('end')
            widget.configure(state='disabled')

    # ── Шаг 0: файлы ─────────────────────────────────────────────────────

    def _on_dnd_drop(self, event):
        paths = []
        for part in re.findall(r'\{([^}]+)\}|(\S+)', event.data):
            p = part[0] or part[1]
            if p:
                paths.append(p)
        for p in paths:
            if p.lower().endswith(('.docx', '.doc')):
                self._add_entry(p)

    def _add_files(self):
        initial = USTAVKI_EXEC_BASE if os.path.isdir(USTAVKI_EXEC_BASE) else os.path.expanduser('~')
        files = filedialog.askopenfilenames(
            title="Выберите таблицы уставок",
            initialdir=initial,
            filetypes=[("Word файлы", "*.docx *.doc"), ("Все файлы", "*.*")],
            parent=self,
        )
        for f in files:
            self._add_entry(f)

    def _add_entry(self, path: str):
        for e in self.ustavki_entries:
            if e['file_path'] == path:
                return
        entry = dict(EMPTY_USTAVKI_ENTRY)
        entry['file_path'] = path
        # Применить данные письма если есть
        entry['letter_num']  = self._lu_letter.get()
        entry['letter_date'] = self._lu_date.get()
        self.ustavki_entries.append(entry)
        self._refresh_lb()

    def _remove_selected(self):
        sel = list(self._lb.curselection())
        for idx in reversed(sel):
            self._lb.delete(idx)
            if idx < len(self.ustavki_entries):
                self.ustavki_entries.pop(idx)
        self._refresh_tree()

    def _refresh_lb(self):
        self._lb.delete(0, 'end')
        for e in self.ustavki_entries:
            self._lb.insert('end', os.path.basename(e['file_path']))

    def _go_parse(self):
        self._step_nb.select(1)
        self._parse_all()

    # ── Шаг 1: парсинг ───────────────────────────────────────────────────

    def _parse_all(self):
        if not HAS_DOCX:
            messagebox.showerror("Ошибка", "python-docx не установлен.", parent=self)
            return
        for entry in self.ustavki_entries:
            try:
                parsed = parse_ustavki_table(entry['file_path'])
                entry.update({
                    'form_type':       parsed['form_type'],
                    'object_name':     parsed['object_name'],
                    'dispatch_name':   parsed['dispatch_name'],
                    'table_number':    parsed['table_number'],
                    'outgoing_letter': parsed['outgoing_letter'],
                    'issue_reason':    parsed.get('issue_reason', ''),
                    'status':          'спарсено',
                })
            except Exception as exc:
                entry['status'] = f'ошибка: {exc}'
        self._apply_letter()
        self._refresh_tree()

    def _apply_letter(self):
        lnum  = self._lu_letter.get()
        ldate = self._lu_date.get()
        lvx   = self._lu_vx.get()
        for entry in self.ustavki_entries:
            entry['letter_num']  = lnum
            entry['letter_date'] = ldate
        self._refresh_tree()

    def _refresh_tree(self):
        self._tree.delete(*self._tree.get_children())
        for entry in self.ustavki_entries:
            self._tree.insert('', 'end', values=(
                os.path.basename(entry.get('file_path','')),
                entry.get('form_type',''),
                entry.get('object_name',''),
                entry.get('dispatch_name',''),
                entry.get('table_number',''),
                entry.get('outgoing_letter',''),
                entry.get('letter_num',''),
                entry.get('letter_date',''),
                entry.get('status',''),
            ))

    def _clear_all(self):
        self.ustavki_entries.clear()
        self._refresh_lb()
        self._refresh_tree()

    # Inline-редактор ячеек
    _TREE_KEYS = [None,'form_type','object_name','dispatch_name',
                  'table_number','outgoing_letter','letter_num','letter_date','status']

    def _on_tree_dclick(self, event):
        tree = self._tree
        region = tree.identify_region(event.x, event.y)
        if region != 'cell':
            return
        col_id = tree.identify_column(event.x)
        row_id = tree.identify_row(event.y)
        if not row_id:
            return
        col_idx = int(col_id.lstrip('#')) - 1
        if col_idx == 0:
            return
        all_ids = list(tree.get_children())
        try:
            ei = all_ids.index(row_id)
        except ValueError:
            return
        if ei >= len(self.ustavki_entries):
            return
        entry = self.ustavki_entries[ei]
        key = self._TREE_KEYS[col_idx] if col_idx < len(self._TREE_KEYS) else None
        if not key:
            return
        x, y, w, h = tree.bbox(row_id, col_id)
        if not w:
            return
        var = tk.StringVar(value=entry.get(key, ''))
        ent = ttk.Entry(tree, textvariable=var)
        ent.place(x=x, y=y, width=w, height=h)
        ent.focus_set(); ent.select_range(0, 'end')

        def _commit(_=None):
            entry[key] = var.get()
            vals = list(tree.item(row_id, 'values'))
            vals[col_idx] = var.get()
            tree.item(row_id, values=vals)
            ent.destroy()

        ent.bind('<Return>', _commit)
        ent.bind('<Escape>', lambda _: ent.destroy())
        ent.bind('<FocusOut>', _commit)

    # ── Шаг 2: запись в таблицы ──────────────────────────────────────────

    def _write_issued_all(self):
        key = '2_write_issued'
        if not HAS_DOCX:
            self._log(key, "ОШИБКА: python-docx не установлен"); return
        vx    = self._lu_vx.get()
        lnum  = self._lu_letter.get()
        ldate = self._lu_date.get()
        llink = self._lu_link.get()

        if not lnum or not ldate:
            messagebox.showwarning("Внимание",
                "Заполните № письма и дату в верхней полоске.", parent=self)
            return

        # Подтверждение
        preview = '\n'.join(
            f"  {os.path.basename(e['file_path'])} → «{lnum} вх-{vx} от {ldate}»"
            for e in self.ustavki_entries
        )
        msg = (f"Запишет в поле «Уставки выставлены» каждого файла:\n\n{preview}\n\n"
               f"Продолжить?")
        if not messagebox.askyesno("Подтверждение", msg, parent=self):
            return

        for entry in self.ustavki_entries:
            fpath = entry.get('file_path', '')
            fname = os.path.basename(fpath)
            try:
                ok = write_issued_to_doc(fpath, lnum, vx, ldate, llink)
                if ok:
                    entry['status'] = 'выставлено'
                    self._log(key, f"OK  {fname}")
                else:
                    entry['status'] = 'нет поля'
                    self._log(key, f"НЕТ ПОЛЯ 'Уставки выставлены'  {fname}")
            except Exception as exc:
                entry['status'] = 'ошибка'
                self._log(key, f"ОШИБКА  {fname}: {exc}")
        self._refresh_tree()

    # ── Шаг 3: реестры ───────────────────────────────────────────────────

    def _write_registries_all(self):
        key = '3_registries'
        if not HAS_OPENPYXL:
            self._log(key, "ОШИБКА: openpyxl не установлен"); return
        vx    = self._lu_vx.get()
        lnum  = self._lu_letter.get()
        ldate = self._lu_date.get()

        # Подтверждение
        forms = [e.get('form_type','?') for e in self.ustavki_entries]
        msg = (f"Записать в реестры?\n\n"
               f"  Новых форм (Реестр.xlsx): {forms.count('new')}\n"
               f"  Старых форм (Регистрация.xlsx): {forms.count('old')}\n"
               f"  Входящее письмо: {lnum} вх-{vx} от {ldate}\n\nПродолжить?")
        if not messagebox.askyesno("Подтверждение", msg, parent=self):
            return

        for entry in self.ustavki_entries:
            fname = os.path.basename(entry.get('file_path', ''))
            form  = entry.get('form_type', '')
            try:
                if form == 'new':
                    found_row, candidates = write_to_registry_new(entry, lnum, vx, ldate)
                    if found_row:
                        entry['registry_row'] = found_row
                        entry['status'] = 'реестр OK'
                        self._log(key, f"OK (стр {found_row})  {fname}")
                    else:
                        self._log(key, f"НЕ НАЙДЕНО  {fname}")
                        dlg = _RegistryNotFoundDialog(
                            self,
                            f"{entry.get('object_name','')} / {entry.get('dispatch_name','')}",
                            candidates,
                        )
                        if dlg.result_row:
                            write_registry_row_manual(dlg.result_row, lnum, vx, ldate)
                            entry['registry_row'] = dlg.result_row
                            entry['status'] = 'реестр OK (ручной)'
                            self._log(key, f"Записано вручную стр {dlg.result_row}")
                        else:
                            entry['status'] = 'пропущено'
                            self._log(key, "Пропущено")
                elif form == 'old':
                    found_row, sheet = write_to_summary_old(entry, lnum, vx)
                    if found_row:
                        entry['registry_row'] = found_row
                        entry['status'] = 'сводная OK'
                        self._log(key, f"OK лист={sheet} стр={found_row}  {fname}")
                    else:
                        entry['status'] = 'не найдено'
                        self._log(key, f"НЕ НАЙДЕНО в сводной  {fname}")
                else:
                    self._log(key, f"Форма неизвестна — пропуск: {fname}")
            except Exception as exc:
                entry['status'] = 'ошибка реестра'
                self._log(key, f"ОШИБКА  {fname}: {exc}")
        self._refresh_tree()

    # ── Шаг 4: раскладка ─────────────────────────────────────────────────

    def _fill_exec_dirs(self, entry: dict):
        """
        Заполняет _current_dir/_archive_dir.

        Структура папок:
          Новый файл:  ...\\Таблицы для исполнения РЗА\\ОБЪЕКТ\\файл.docx
          Старый файл: ...\\Таблицы уставок РЗА\\ОБЪЕКТ\\файл.docx  (кандидат на архив)

        После операции:
          Старый → ...\\Архив таблиц РЗА\\ОБЪЕКТ\\        (_archive_dir)
          Новый  → ...\\Таблицы уставок РЗА\\ОБЪЕКТ\\     (_current_dir)
        """
        short = get_object_short_name_from_path(entry['file_path'])
        if not short:
            obj = entry.get('object_name', '')
            short = match_object_to_short_name(obj) if obj else ''
        if not short:
            short = os.path.basename(os.path.dirname(entry['file_path']))
        entry['short_name'] = short

        if short:
            current_dir = os.path.join(USTAVKI_ARCHIVE_BASE, short)       # Таблицы уставок РЗА\ОБЪЕКТ
            archive_dir = os.path.join(USTAVKI_REAL_ARCHIVE_BASE, short)   # Архив таблиц РЗА\ОБЪЕКТ
        else:
            current_dir = archive_dir = ''

        entry['_current_dir'] = current_dir
        entry['_archive_dir'] = archive_dir
        return short

    def _find_candidates_a(self):
        """Вариант А: кандидат по схожести ИМЕНИ ФАЙЛА в Таблицы уставок РЗА\\Объект."""
        self._candidates_found = True
        self._arc_tree.delete(*self._arc_tree.get_children())
        for entry in self.ustavki_entries:
            short = self._fill_exec_dirs(entry)
            candidates = find_archive_candidates_by_filename(entry['file_path'])
            top = candidates[0] if candidates else ('', '', 0.0)
            entry['archive_candidate'] = top[0]
            self._arc_tree.insert('', 'end', values=(
                os.path.basename(entry['file_path']),
                short,
                top[1] if top else '',
                f"{top[2]:.2f}" if len(top) > 2 else '',
                'А',
            ))

    def _find_candidates_b(self):
        """Вариант Б: кандидат по схожести ДИСПЕТЧЕРСКОГО НАИМЕНОВАНИЯ."""
        self._arc_tree.delete(*self._arc_tree.get_children())
        for entry in self.ustavki_entries:
            short = self._fill_exec_dirs(entry)
            dispatch = entry.get('dispatch_name', '') or entry.get('object_name', '')
            candidates = find_archive_candidates_by_dispatch(entry['file_path'], dispatch)
            top = candidates[0] if candidates else ('', '', 0.0, '')
            entry['archive_candidate'] = top[0]
            self._arc_tree.insert('', 'end', values=(
                os.path.basename(entry['file_path']),
                short,
                top[1] if top else '',
                f"{top[2]:.2f}" if len(top) > 2 else '',
                'Б',
            ))

    def _on_arc_dclick(self, event):
        tree = self._arc_tree
        row_id = tree.identify_row(event.y)
        if not row_id:
            return
        all_ids = list(tree.get_children())
        try:
            idx = all_ids.index(row_id)
        except ValueError:
            return
        if idx >= len(self.ustavki_entries):
            return
        entry = self.ustavki_entries[idx]

        # Собираем кандидатов из обоих источников
        cands_a = find_archive_candidates_by_filename(entry['file_path'], top_n=10)
        dispatch = entry.get('dispatch_name', '') or entry.get('object_name', '')
        cands_b = find_archive_candidates_by_dispatch(entry['file_path'], dispatch, top_n=10) if HAS_DOCX else []

        # Объединяем: (path, name, score, method)
        seen = set()
        all_cands = []
        for path, name, score in cands_a:
            if path not in seen:
                seen.add(path)
                all_cands.append((path, name, score, 'А'))
        for path, name, score, _ in cands_b:
            if path not in seen:
                seen.add(path)
                all_cands.append((path, name, score, 'Б'))
        all_cands.sort(key=lambda x: x[2], reverse=True)

        if not all_cands:
            messagebox.showinfo("Кандидаты", "Похожих файлов не найдено.", parent=self)
            return

        dlg = tk.Toplevel(self)
        dlg.title("Выбор архивной таблицы")
        dlg.grab_set()
        ttk.Label(dlg,
                  text=f"Кандидаты для:\n{os.path.basename(entry['file_path'])}",
                  wraplength=460, font=('', '10', 'bold')).pack(padx=12, pady=(10, 4))
        cols2 = ('name', 'score', 'method')
        tv = ttk.Treeview(dlg, columns=cols2, show='headings', height=10)
        tv.heading('name',   text='Имя файла');   tv.column('name',   width=320)
        tv.heading('score',  text='Схожесть');    tv.column('score',  width=70)
        tv.heading('method', text='Метод');       tv.column('method', width=50)
        for cpath, cname, cscore, method in all_cands:
            tv.insert('', 'end', iid=cpath, values=(cname, f"{cscore:.2f}", method))
        tv.pack(padx=12, fill='both', expand=True)

        def _pick():
            sel = tv.selection()
            if sel:
                entry['archive_candidate'] = sel[0]
                vals = list(tree.item(row_id, 'values'))
                vals[2] = os.path.basename(sel[0])
                tree.item(row_id, values=vals)
            dlg.destroy()

        ttk.Button(dlg, text="Выбрать", command=_pick).pack(pady=8)
        dlg.transient(self)
        dlg.wait_window()

    def _move_files(self):
        # Проверка: была ли нажата кнопка «Найти кандидатов»
        if not self._candidates_found:
            messagebox.showwarning("Нет данных",
                "Сначала нажмите «Найти кандидатов на архив».", parent=self)
            return

        if not self.ustavki_entries:
            messagebox.showwarning("Нет файлов",
                "Список записей пуст. Добавьте файлы на шаге 0.", parent=self)
            return

        # Строим список операций для подтверждения
        lines = []
        for entry in self.ustavki_entries:
            cd  = entry.get('_current_dir', '')
            ad  = entry.get('_archive_dir', '')
            arc = entry.get('archive_candidate', '')
            fn  = os.path.basename(entry['file_path'])
            if cd and ad:
                lines.append(
                    f"  {fn}\n    → Текущие: {cd}\n"
                    + (f"    архивировать: {os.path.basename(arc)}\n    → Архив: {ad}\n" if arc else '')
                )
            else:
                lines.append(f"  {fn}  [папки объекта не найдены — будет пропущено]")
        msg = "Разложить файлы?\n\n" + '\n'.join(lines[:10])
        if len(lines) > 10:
            msg += f"\n...и ещё {len(lines)-10} файл(ов)"
        if not messagebox.askyesno("Подтверждение", msg, parent=self):
            return

        moved = errors = 0
        for entry in self.ustavki_entries:
            cd = entry.get('_current_dir', '')
            ad = entry.get('_archive_dir', '')
            if not cd or not ad:
                continue
            try:
                new_path = move_table_files(entry, ad, cd)
                if new_path:
                    entry['current_path'] = new_path
                    entry['status'] = 'разложено'
                    moved += 1
            except Exception as exc:
                errors += 1
                messagebox.showerror("Ошибка раскладки",
                    f"{os.path.basename(entry['file_path'])}: {exc}", parent=self)
        self._refresh_tree()
        messagebox.showinfo("Раскладка завершена",
            f"Разложено: {moved}  Ошибок: {errors}", parent=self)
        self._save_current_session()

    # ── Шаг 5: изменения (синие строки) ──────────────────────────────────

    def _build_step5(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)
        ttk.Label(parent,
            text="Ищет строки синего цвета (изменения) во всех таблицах .docx "
                 "и формирует сводный отчёт Word.\n"
                 "Файлы берутся из текущего списка (поле file_path каждой записи).",
            wraplength=700, foreground='gray').grid(row=0, column=0, sticky='w', pady=(0, 6))
        txt = tk.Text(parent, height=14, wrap='word', state='disabled', font=('Consolas', 9))
        sb = ttk.Scrollbar(parent, orient='vertical', command=txt.yview)
        txt.configure(yscrollcommand=sb.set)
        txt.grid(row=1, column=0, sticky='nsew')
        sb.grid(row=1, column=1, sticky='ns')
        self._logs['5'] = txt
        bf = ttk.Frame(parent)
        bf.grid(row=2, column=0, pady=(6, 0), sticky='w')
        ttk.Button(bf, text="Создать отчёт изменений →",
                   command=self._create_changes_report).pack(side='left', padx=4)

    def _create_changes_report(self):
        if not HAS_DOCX:
            self._log('5', "ОШИБКА: python-docx не установлен"); return
        if not self.ustavki_entries:
            self._log('5', "Список записей пуст. Добавьте файлы на шаге 0."); return
        default_name = f"Изменения_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        out_path = filedialog.asksaveasfilename(
            title="Сохранить отчёт изменений",
            initialdir=os.path.expanduser('~'),
            initialfile=default_name,
            defaultextension='.docx',
            filetypes=[("Word файлы", "*.docx"), ("Все файлы", "*.*")],
            parent=self,
        )
        if not out_path:
            return
        msg = (f"Создать отчёт изменений для {len(self.ustavki_entries)} таблиц?\n\n"
               f"Сохранить как:\n{out_path}\n\nПродолжить?")
        if not messagebox.askyesno("Подтверждение", msg, parent=self):
            return
        self._log('5', f"Анализ {len(self.ustavki_entries)} файлов...")
        try:
            generate_changes_report(self.ustavki_entries, out_path)
            self._log('5', f"Отчёт создан: {out_path}")
        except Exception as exc:
            self._log('5', f"ОШИБКА: {exc}")

    # ── Шаг 6: карты Visio ───────────────────────────────────────────────

    def _build_step6(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(3, weight=1)

        cfg = ttk.LabelFrame(parent, text="Пути к папкам", padding=6)
        cfg.grid(row=0, column=0, sticky='ew', pady=(0, 4))
        cfg.columnconfigure(1, weight=1)
        ttk.Label(cfg, text="Папка карт Visio:").grid(row=0, column=0, sticky='e', padx=(0, 6))
        self._maps_folder_var = tk.StringVar(value=MAPS_FOLDER)
        ttk.Entry(cfg, textvariable=self._maps_folder_var, width=50).grid(row=0, column=1, sticky='ew')
        ttk.Button(cfg, text="…", command=lambda: self._browse_folder(self._maps_folder_var),
                   width=3).grid(row=0, column=2)
        ttk.Label(cfg, text="Папка PDF:").grid(row=1, column=0, sticky='e', padx=(0, 6))
        self._pdf_folder_var = tk.StringVar(value=MAPS_PDF_FOLDER)
        ttk.Entry(cfg, textvariable=self._pdf_folder_var, width=50).grid(row=1, column=1, sticky='ew')
        ttk.Button(cfg, text="…", command=lambda: self._browse_folder(self._pdf_folder_var),
                   width=3).grid(row=1, column=2)

        inp = ttk.LabelFrame(parent,
            text="Ручной ввод (заполняется автоматически из сессии или вручную)", padding=6)
        inp.grid(row=1, column=0, sticky='ew', pady=(0, 4))
        inp.columnconfigure(1, weight=1)
        ttk.Label(inp, text="Старый файл таблицы\n(гиперссылка для замены):",
                  justify='right').grid(row=0, column=0, sticky='e', padx=(0, 6), pady=3)
        self._old_table_path_var = tk.StringVar()
        old_row = ttk.Frame(inp); old_row.grid(row=0, column=1, sticky='ew', pady=3)
        old_row.columnconfigure(0, weight=1)
        ttk.Entry(old_row, textvariable=self._old_table_path_var, width=52).grid(row=0, column=0, sticky='ew')
        ttk.Button(old_row, text="…", command=lambda: self._browse_file(self._old_table_path_var),
                   width=3).grid(row=0, column=1, padx=(4, 0))
        ttk.Label(inp, text="Новый файл таблицы\n(будет вставлен как ссылка):",
                  justify='right').grid(row=1, column=0, sticky='e', padx=(0, 6), pady=3)
        self._new_table_path_var = tk.StringVar()
        new_row = ttk.Frame(inp); new_row.grid(row=1, column=1, sticky='ew', pady=3)
        new_row.columnconfigure(0, weight=1)
        ttk.Entry(new_row, textvariable=self._new_table_path_var, width=52).grid(row=0, column=0, sticky='ew')
        ttk.Button(new_row, text="…", command=lambda: self._browse_file(self._new_table_path_var),
                   width=3).grid(row=0, column=1, padx=(4, 0))
        ttk.Label(inp, text="(при наличии сессии — новый путь берётся из current_path каждой записи)",
                  foreground='gray').grid(row=2, column=0, columnspan=2, sticky='w', pady=(0, 2))
        ttk.Label(inp, text="Объект (краткое имя\n= имя папки):").grid(
            row=3, column=0, sticky='e', padx=(0, 6), pady=3)
        obj_row = ttk.Frame(inp); obj_row.grid(row=3, column=1, sticky='ew', pady=3)
        obj_row.columnconfigure(0, weight=1)
        self._manual_object_var = tk.StringVar()
        ttk.Entry(obj_row, textvariable=self._manual_object_var, width=30).grid(row=0, column=0, sticky='w')
        ttk.Button(obj_row, text="Определить из пути нового файла",
                   command=self._auto_detect_object).grid(row=0, column=1, padx=(8, 0))

        ttk.Label(parent,
            text="Для каждой таблицы из списка:\n"
                 "  • Находит .vsdx файл по краткому имени объекта\n"
                 "  • Заменяет гиперссылку (старый → новый путь к .docx)\n"
                 "  • Сохраняет .vsdx и экспортирует PDF\n"
                 "ТРЕБУЕТ: Microsoft Visio и pywin32",
            wraplength=700, foreground='gray').grid(row=2, column=0, sticky='w', pady=(0, 4))
        txt = tk.Text(parent, height=10, wrap='word', state='disabled', font=('Consolas', 9))
        sb = ttk.Scrollbar(parent, orient='vertical', command=txt.yview)
        txt.configure(yscrollcommand=sb.set)
        txt.grid(row=3, column=0, sticky='nsew')
        sb.grid(row=3, column=1, sticky='ns')
        self._logs['6'] = txt
        bf = ttk.Frame(parent)
        bf.grid(row=4, column=0, pady=(6, 0), sticky='w')
        ttk.Button(bf, text="Обновить карты Visio (из сессии) →",
                   command=self._update_maps_all).pack(side='left', padx=4)
        ttk.Button(bf, text="Обновить карту для одного объекта →",
                   command=self._update_single_map).pack(side='left', padx=4)

    def _browse_folder(self, var: tk.StringVar):
        d = filedialog.askdirectory(initialdir=var.get() or os.path.expanduser('~'), parent=self)
        if d:
            var.set(d)

    def _browse_file(self, var: tk.StringVar):
        initial = os.path.dirname(var.get()) if var.get() else os.path.expanduser('~')
        f = filedialog.askopenfilename(
            initialdir=initial,
            filetypes=[("Word файлы", "*.docx *.doc"), ("Все файлы", "*.*")],
            parent=self,
        )
        if f:
            var.set(f.replace('/', '\\'))

    def _auto_detect_object(self):
        new_path = self._new_table_path_var.get().strip()
        if not new_path:
            messagebox.showwarning("Нет пути", "Сначала выберите новый файл таблицы.", parent=self)
            return
        short = get_object_short_name_from_path(new_path)
        if short:
            self._manual_object_var.set(short)
        else:
            short = os.path.basename(os.path.dirname(new_path))
            self._manual_object_var.set(short)
            messagebox.showinfo("Объект",
                f"Объект не найден в справочнике.\nУстановлено: {short}", parent=self)

    def _update_maps_all(self):
        maps_folder = self._maps_folder_var.get()
        if not self.ustavki_entries:
            self._log('6', "Список записей пуст. Добавьте файлы на шаге 0."); return
        objects_to_update = [
            match_object_to_short_name(e.get('object_name', ''))
            for e in self.ustavki_entries
            if match_object_to_short_name(e.get('object_name', ''))
        ]
        msg = (f"Обновить карты Visio для {len(objects_to_update)} объектов?\n\n"
               f"Папка карт: {maps_folder}\n\n"
               + '\n'.join(f"  • {o}" for o in objects_to_update[:10])
               + (f"\n  ...и ещё {len(objects_to_update)-10}" if len(objects_to_update) > 10 else '')
               + "\n\nЭто откроет Microsoft Visio для каждого объекта!\nПродолжить?")
        if not messagebox.askyesno("Подтверждение", msg, parent=self):
            return
        pdf_folder = self._pdf_folder_var.get()
        for entry in self.ustavki_entries:
            short = match_object_to_short_name(entry.get('object_name', ''))
            if not short:
                self._log('6', f"Объект не распознан: {entry.get('object_name','')}"); continue
            visio_path = os.path.join(maps_folder, short + '.vsdx')
            if not os.path.exists(visio_path):
                visio_path = os.path.join(maps_folder, short + '.vsd')
            if not os.path.exists(visio_path):
                self._log('6', f"Visio не найден: {short}.vsdx / .vsd"); continue
            old_path  = entry.get('archive_candidate', '')
            new_path  = entry.get('current_path', entry.get('file_path', ''))
            table_num = entry.get('table_number', '')
            if not old_path:
                self._log('6', f"  ПРОПУСК {short}: не найден предыдущий файл (archive_candidate пуст)"); continue
            self._log('6', f"Обновляю: {short}  {os.path.basename(visio_path)}")
            ok, msg_r = update_visio_map(visio_path, old_path, new_path, table_num, pdf_folder)
            self._log('6', f"  {'OK' if ok else 'ERR'}  {msg_r}")
        self._save_current_session()
        self._log('6', "--- Завершено ---")

    def _update_single_map(self):
        old_path = self._old_table_path_var.get().strip()
        new_path = self._new_table_path_var.get().strip()
        short    = self._manual_object_var.get().strip()
        if not new_path:
            messagebox.showwarning("Нет данных", "Укажите новый файл таблицы.", parent=self); return
        if not short:
            messagebox.showwarning("Нет данных",
                "Укажите краткое имя объекта (или нажмите «Определить»).", parent=self); return
        maps_folder = self._maps_folder_var.get()
        visio_path = os.path.join(maps_folder, short + '.vsdx')
        if not os.path.exists(visio_path):
            visio_path = os.path.join(maps_folder, short + '.vsd')
        if not os.path.exists(visio_path):
            self._log('6', f"Visio не найден: {short}.vsdx / .vsd в {maps_folder}"); return
        msg = (f"Обновить карту?\n\n"
               f"  Объект:  {short}\n"
               f"  Visio:   {os.path.basename(visio_path)}\n"
               f"  Старый:  {os.path.basename(old_path) if old_path else '(не указан)'}\n"
               f"  Новый:   {os.path.basename(new_path)}\n\n"
               f"Откроется Microsoft Visio. Продолжить?")
        if not old_path:
            messagebox.showwarning("Нет данных",
                "Укажите старый файл таблицы (гиперссылку для замены).\n"
                "Без него невозможно определить, какую ссылку заменить.", parent=self)
            return
        if not messagebox.askyesno("Подтверждение", msg, parent=self):
            return
        pdf_folder = self._pdf_folder_var.get()
        self._log('6', f"Обновляю: {short}  {os.path.basename(visio_path)}")
        ok, result_msg = update_visio_map(visio_path, old_path, new_path, '', pdf_folder)
        self._log('6', f"  {'OK' if ok else 'ERR'}  {result_msg}")
        self._save_current_session()

    # ── Шаг 7: ДЭБ ───────────────────────────────────────────────────────

    def _build_step7(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(2, weight=1)

        # Настройки
        cfg = ttk.LabelFrame(parent, text="Настройки браузера", padding=6)
        cfg.grid(row=0, column=0, sticky='ew', pady=(0, 4))
        cfg.columnconfigure(1, weight=1)

        ttk.Label(cfg, text="yandexdriver.exe:").grid(row=0, column=0, sticky='e', padx=(0, 6))
        self._deb_driver_var = tk.StringVar()
        # По умолчанию ищем рядом с exe
        default_drv = os.path.join(
            os.path.dirname(sys.executable) if getattr(sys, 'frozen', False)
            else os.path.dirname(os.path.abspath(__file__)),
            'yandexdriver.exe'
        )
        if os.path.exists(default_drv):
            self._deb_driver_var.set(default_drv)
        dr_row = ttk.Frame(cfg); dr_row.grid(row=0, column=1, sticky='ew')
        dr_row.columnconfigure(0, weight=1)
        ttk.Entry(dr_row, textvariable=self._deb_driver_var, width=52).grid(row=0, column=0, sticky='ew')
        ttk.Button(dr_row, text="…",
                   command=lambda: self._browse_driver_exe(self._deb_driver_var),
                   width=3).grid(row=0, column=1, padx=(4, 0))
        ttk.Label(cfg,
                  text="Скачать yandexdriver: https://yandex.ru/dev/yandexdriver/\n"
                       "Поместите yandexdriver.exe рядом с exe (или укажите путь выше).",
                  foreground='gray').grid(row=1, column=0, columnspan=2, sticky='w', pady=(2, 0))

        ttk.Label(parent,
                  text="Загружает обновлённые карты уставок (Visio + PDF) в ДЭБ через Яндекс Браузер.\n"
                       "Для каждой записи: берёт dispatch_name → ищет в каталоге ДЭБ → "
                       "заменяет файлы → сохраняет.\n"
                       "Требует: selenium, pyautogui, yandexdriver.exe рядом с exe.",
                  wraplength=720, foreground='gray').grid(row=1, column=0, sticky='w', pady=(0, 4))

        txt = tk.Text(parent, height=12, wrap='word', state='disabled', font=('Consolas', 9))
        sb = ttk.Scrollbar(parent, orient='vertical', command=txt.yview)
        txt.configure(yscrollcommand=sb.set)
        txt.grid(row=2, column=0, sticky='nsew')
        sb.grid(row=2, column=1, sticky='ns')
        self._logs['7'] = txt

        bf = ttk.Frame(parent)
        bf.grid(row=3, column=0, pady=(6, 0), sticky='w')
        ttk.Button(bf, text="Загрузить в ДЭБ (все записи) →",
                   command=self._upload_deb_all).pack(side='left', padx=4)
        ttk.Button(bf, text="Открыть каталог ДЭБ в браузере",
                   command=self._open_deb_url).pack(side='left', padx=4)

    def _browse_driver_exe(self, var: tk.StringVar):
        f = filedialog.askopenfilename(
            title="Выберите yandexdriver.exe / chromedriver.exe",
            initialdir=os.path.dirname(var.get()) if var.get() else os.path.expanduser('~'),
            filetypes=[("Exe файлы", "*.exe"), ("Все файлы", "*.*")],
            parent=self,
        )
        if f:
            var.set(f.replace('/', '\\'))

    def _open_deb_url(self):
        import webbrowser
        webbrowser.open(DEB_MAPS_URL)

    def _upload_deb_all(self):
        if not HAS_SELENIUM:
            self._log('7', "ОШИБКА: selenium не установлен  →  pip install selenium"); return
        if not HAS_PYAUTOGUI:
            self._log('7', "ОШИБКА: pyautogui не установлен  →  pip install pyautogui"); return
        if not self.ustavki_entries:
            self._log('7', "Список записей пуст. Добавьте файлы на шаге 0."); return

        driver_path = self._deb_driver_var.get().strip()
        maps_folder = self._maps_folder_var.get() if hasattr(self, '_maps_folder_var') else MAPS_FOLDER
        pdf_folder  = self._pdf_folder_var.get()  if hasattr(self, '_pdf_folder_var')  else MAPS_PDF_FOLDER

        objects_info = []
        for entry in self.ustavki_entries:
            short = match_object_to_short_name(entry.get('object_name', ''))
            if not short:
                continue
            dispatch = entry.get('dispatch_name', '') or entry.get('object_name', '')
            visio_path = ''
            for ext in ('.vsdx', '.vsd'):
                cand = os.path.join(maps_folder, short + ext)
                if os.path.exists(cand):
                    visio_path = cand
                    break
            pdf_path = os.path.join(pdf_folder, short + '.pdf')
            objects_info.append((dispatch, visio_path, pdf_path))

        if not objects_info:
            self._log('7', "Нет объектов с распознанным коротким именем. "
                           "Проверьте шаг 1 (парсировать данные).")
            return

        msg = (f"Загрузить в ДЭБ для {len(objects_info)} объектов?\n\n" +
               '\n'.join(f"  • {d}" + (f"\n    Visio: {os.path.basename(v)}" if v else "  [Visio не найден]")
                         for d, v, _ in objects_info[:8]) +
               (f"\n  ...и ещё {len(objects_info)-8}" if len(objects_info) > 8 else '') +
               "\n\nОткроется Яндекс Браузер. Продолжить?")
        if not messagebox.askyesno("Подтверждение", msg, parent=self):
            return

        self._log('7', f"Начало загрузки {len(objects_info)} объектов...")
        for dispatch, visio_path, pdf_path in objects_info:
            self._log('7', f"→ {dispatch}")
            ok, msg_r = upload_to_deb_entry(dispatch, visio_path, pdf_path, driver_path)
            self._log('7', f"  {'OK' if ok else 'ERR'}  {msg_r}")
        self._log('7', "--- Завершено ---")

    # ── Сессия ────────────────────────────────────────────────────────────

    def _load_session(self):
        session = load_session()
        if 'letter' in session:
            d = session['letter'].get('in_data', {})
            self._lu_vx.set(d.get('incoming_num',''))
            self._lu_letter.set(d.get('letter_num',''))
            self._lu_date.set(d.get('date',''))
            self._lu_link.set(d.get('file_link',''))
            self.in_data = d
        if 'ustavki' in session:
            self.ustavki_entries = session['ustavki']
            self._refresh_lb()
            self._refresh_tree()

    def _save_current_session(self):
        self.in_data = {
            'incoming_num': self._lu_vx.get(),
            'letter_num':   self._lu_letter.get(),
            'date':         self._lu_date.get(),
            'file_link':    self._lu_link.get(),
        }
        session = load_session()
        session['ustavki'] = self.ustavki_entries
        save_session(session)

    # ── Центровка ─────────────────────────────────────────────────────────

    def _center_window(self):
        self.update_idletasks()
        w, h = self.winfo_width(), self.winfo_height()
        sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
        self.geometry(f"{max(w,1000)}x{max(h,650)}+{(sw-max(w,1000))//2}+{(sh-max(h,650))//2}")


# ── Точка входа ───────────────────────────────────────────────────────────────

def main():
    if not HAS_DOCX:
        root = tk.Tk(); root.withdraw()
        messagebox.showwarning("python-docx не установлен",
            "Парсинг и запись в .docx недоступны.\npip install python-docx")
        root.destroy()
    app = UstavkiFoldersApp()
    app.mainloop()


if __name__ == '__main__':
    main()
